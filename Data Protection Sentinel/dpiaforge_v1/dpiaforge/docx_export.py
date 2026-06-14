"""
DPIAforge — Word document export using python-docx.
"""
import io
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colours
DEEP_BLUE  = RGBColor(0x1E, 0x3A, 0x8A)
SLATE_GRAY = RGBColor(0x33, 0x41, 0x55)
CYAN       = RGBColor(0x22, 0xD3, 0xEE)
INDIGO     = RGBColor(0x4F, 0x46, 0xE5)
LIGHT_BG   = "EFF6FF"
HEADER_BG  = "1E3A8A"
ALT_ROW    = "F0F9FF"

REGULATION_FULL = {
    "GDPR":                "EU General Data Protection Regulation 2016/679",
    "Zimbabwe CDPA":       "Zimbabwe Cyber and Data Protection Act [Chapter 12:07]",
    "South Africa POPIA":  "Protection of Personal Information Act 4 of 2013",
    "UAE PDPL":            "UAE Federal Decree-Law No. 45 of 2021",
    "Saudi PDPL":          "Saudi Arabia Personal Data Protection Law",
    "Qatar DPL":           "Qatar Personal Data Privacy Protection Law No. 13 of 2016",
}


def _shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc, text: str, level: int = 1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color or (DEEP_BLUE if level == 1 else SLATE_GRAY)
        run.font.name = "Calibri"
    return h


def _info_table(doc, rows: list[tuple], label_width=2.2, value_width=4.7):
    """2-column label/value info table."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    lw = int(label_width * 1440)
    vw = int(value_width * 1440)
    for i, (label, value) in enumerate(rows):
        bg = ALT_ROW if i % 2 else "FFFFFF"
        lc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        # widths
        lc.width = lw; vc.width = vw
        _shading(lc, LIGHT_BG)
        _shading(vc, bg)
        # content
        lp = lc.paragraphs[0]
        lr = lp.add_run(str(label))
        lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = DEEP_BLUE
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value) if value else "—")
        vr.font.size = Pt(9.5)
    doc.add_paragraph()
    return tbl


def _risk_table(doc, risks: list):
    if not risks:
        doc.add_paragraph("No risks recorded.", style="Normal")
        return
    headers = ["Risk Description", "Likelihood", "Impact", "Mitigation"]
    widths  = [3.2, 1.0, 1.0, 3.7]
    tbl = doc.add_table(rows=1 + len(risks), cols=4)
    tbl.style = "Table Grid"
    # Header row
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.width = int(widths[i] * 1440)
        _shading(c, HEADER_BG)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    LEVEL_COLORS = {"High": "FEE2E2", "Medium": "FEF3C7", "Low": "D1FAE5"}
    for i, risk in enumerate(risks):
        row = tbl.rows[i + 1]
        vals = [
            risk.get("desc", ""),
            risk.get("likelihood", ""),
            risk.get("impact", ""),
            risk.get("mitigation", ""),
        ]
        bg = LEVEL_COLORS.get(risk.get("likelihood", ""), "FFFFFF")
        for j, val in enumerate(vals):
            c = row.cells[j]
            c.width = int(widths[j] * 1440)
            _shading(c, ALT_ROW if j == 0 else bg)
            c.paragraphs[0].add_run(str(val)).font.size = Pt(9)
    doc.add_paragraph()


def _render_ai_text(doc, text: str):
    """Parse markdown-ish AI output into Word paragraphs."""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            for r in p.runs:
                r.font.color.rgb = INDIGO
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for r in p.runs:
                r.font.color.rgb = SLATE_GRAY
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            for r in p.runs:
                r.font.color.rgb = DEEP_BLUE
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.runs[0].font.size = Pt(10.5) if p.runs else None
        else:
            # Inline bold **text**
            p = doc.add_paragraph()
            import re
            parts = re.split(r'\*\*(.+?)\*\*', stripped)
            for k, part in enumerate(parts):
                if not part:
                    continue
                r = p.add_run(part)
                r.bold = (k % 2 == 1)
                r.font.size = Pt(10.5)


def generate_docx(dpia: dict) -> io.BytesIO:
    doc = Document()

    # ── page setup ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── default style ───────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    regulation     = dpia.get("regulation", "GDPR")
    regulation_full = REGULATION_FULL.get(regulation, regulation)
    created        = dpia.get("created_at", datetime.utcnow().strftime("%Y-%m-%d"))

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = brand.add_run("DPIA")
    r1.font.size = Pt(36); r1.font.bold = True; r1.font.color.rgb = DEEP_BLUE; r1.font.name = "Calibri"
    r2 = brand.add_run("forge")
    r2.font.size = Pt(36); r2.font.bold = True; r2.font.color.rgb = CYAN; r2.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("by ALI MOYO")
    sr.font.size = Pt(11); sr.font.color.rgb = SLATE_GRAY; sr.font.italic = True

    doc.add_paragraph()

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.border_bottom = True
    sr2 = sep.add_run("DATA PROTECTION IMPACT ASSESSMENT")
    sr2.font.size = Pt(14); sr2.font.bold = True; sr2.font.color.rgb = SLATE_GRAY

    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(dpia.get("title", "Untitled DPIA"))
    tr.font.size = Pt(22); tr.font.bold = True; tr.font.color.rgb = DEEP_BLUE

    doc.add_paragraph()

    reg_p = doc.add_paragraph()
    reg_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = reg_p.add_run(regulation_full)
    rr.font.size = Pt(12); rr.font.italic = True; rr.font.color.rgb = INDIGO

    doc.add_paragraph()
    doc.add_paragraph()

    _info_table(doc, [
        ("Reference Number",  dpia.get("ref_number")),
        ("Organisation",      dpia.get("org_name")),
        ("Department",        dpia.get("department")),
        ("Data Controller",   dpia.get("controller_name")),
        ("DPO / Privacy Lead",dpia.get("dpo_name")),
        ("DPO Email",         dpia.get("dpo_email")),
        ("Processing Activity", dpia.get("activity_type")),
        ("Status",            (dpia.get("status", "draft")).replace("_", " ").title()),
        ("Date Created",      created[:10]),
        ("Regulation",        regulation),
    ])

    doc.add_paragraph()
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = conf.add_run("CONFIDENTIAL — For authorised personnel only")
    cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8); cr.font.italic = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1 — PROCESSING DETAILS
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, "1. Processing Activity Details")
    cats  = ", ".join(dpia.get("data_categories", [])) or "Not specified"
    scats = ", ".join(dpia.get("special_cats", []))     or "None"
    _info_table(doc, [
        ("Activity Type",       dpia.get("activity_type")),
        ("Description",         dpia.get("activity_desc")),
        ("Purpose(s)",          dpia.get("purpose")),
        ("Legal Basis",         dpia.get("legal_basis")),
        ("Data Categories",     cats),
        ("Special Categories",  scats),
        ("Data Subjects",       dpia.get("data_subjects")),
        ("Estimated Volume",    dpia.get("subject_count")),
        ("Retention Period",    dpia.get("retention")),
    ])

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2 — SYSTEMS & TRANSFERS
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, "2. Systems, Processors & International Transfers")
    _info_table(doc, [
        ("Processing Systems",      dpia.get("systems")),
        ("Third-party Processors",  dpia.get("processors")),
        ("International Transfer",  dpia.get("intl_transfer")),
        ("Destination Countries",   dpia.get("transfer_dest")),
        ("Transfer Mechanism",      dpia.get("transfer_mech")),
    ])

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3 — NECESSITY & PROPORTIONALITY
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, "3. Necessity & Proportionality")
    _heading(doc, "3.1 Necessity Assessment", level=2)
    doc.add_paragraph(dpia.get("necessity") or "To be completed.")
    _heading(doc, "3.2 Proportionality Assessment", level=2)
    doc.add_paragraph(dpia.get("proportionality") or "To be completed.")

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4 — RISK ASSESSMENT
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "4. Risk Assessment")
    _info_table(doc, [
        ("Overall Risk Level", dpia.get("overall_risk")),
        ("Residual Risk Level", dpia.get("residual_risk")),
    ])
    _heading(doc, "4.1 Identified Risks", level=2)
    _risk_table(doc, dpia.get("risks", []))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5 — CONSULTATION
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, "5. Consultation Record")
    _info_table(doc, [
        ("DPO Consulted",                dpia.get("dpo_consulted")),
        ("Supervisory Authority",        dpia.get("auth_consulted")),
        ("Data Subjects Consulted",      dpia.get("subjects_consulted")),
        ("Consultation Notes",           dpia.get("consult_notes")),
    ])

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6 — AI RESEARCH (if present)
    # ═══════════════════════════════════════════════════════════════════════
    if dpia.get("ai_research"):
        doc.add_page_break()
        _heading(doc, "6. AI Research Summary")
        note = doc.add_paragraph()
        nr = note.add_run("AI-assisted research output — review and validate before submission.")
        nr.font.size = Pt(9); nr.font.italic = True; nr.font.color.rgb = INDIGO
        doc.add_paragraph()
        _render_ai_text(doc, dpia["ai_research"])

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 7 — FULL AI-GENERATED DPIA (if present)
    # ═══════════════════════════════════════════════════════════════════════
    if dpia.get("ai_full_dpia"):
        doc.add_page_break()
        _heading(doc, "7. Full DPIA Analysis (AI-Generated)")
        note = doc.add_paragraph()
        nr = note.add_run("AI-generated content — DPO review required before regulatory submission.")
        nr.font.size = Pt(9); nr.font.italic = True; nr.font.color.rgb = INDIGO
        doc.add_paragraph()
        _render_ai_text(doc, dpia["ai_full_dpia"])

    # ═══════════════════════════════════════════════════════════════════════
    # SIGN-OFF PAGE
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, "Sign-Off & Approval")
    _info_table(doc, [
        ("DPO / Privacy Officer", ""),
        ("Name",      dpia.get("dpo_name") or "________________________"),
        ("Signature", "________________________"),
        ("Date",      "________________________"),
        ("", ""),
        ("Senior Management", ""),
        ("Name",      "________________________"),
        ("Title",     "________________________"),
        ("Signature", "________________________"),
        ("Date",      "________________________"),
    ])

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"Generated by DPIAforge · by ALI MOYO · {datetime.utcnow().strftime('%d %B %Y')}")
    fr.font.size = Pt(8.5); fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8); fr.font.italic = True

    # ── serialise ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
