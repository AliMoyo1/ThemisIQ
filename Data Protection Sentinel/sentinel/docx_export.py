"""
Data Protection Sentinel — Word document export.
Updated branding: Data Protection Sentinel by Ali Moyo.
Green theme matching the Donezo light UI.
"""
import io
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colours — green Donezo theme
DARK_GREEN  = RGBColor(0x16, 0x65, 0x34)   # #166534
MID_GREEN   = RGBColor(0x15, 0x80, 0x3D)   # #15803d
LIGHT_GREEN = RGBColor(0x05, 0x96, 0x69)   # #059669
TEXT_DARK   = RGBColor(0x11, 0x18, 0x27)   # #111827
TEXT_GREY   = RGBColor(0x6B, 0x72, 0x80)   # #6b7280
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
RISK_VH     = RGBColor(0xDC, 0x26, 0x26)   # red
RISK_H      = RGBColor(0xD9, 0x77, 0x06)   # amber
RISK_M      = RGBColor(0xCA, 0x8A, 0x04)   # yellow
RISK_L      = RGBColor(0x05, 0x96, 0x69)   # green

# Hex shading strings
HDR_BG      = "166534"  # dark green header
HDR_BG2     = "15803d"  # mid green section header
ACCENT_BG   = "f0fdf4"  # very light green tint
ALT_ROW     = "f9fafb"  # light grey alt row
RISK_VH_BG  = "fee2e2"
RISK_H_BG   = "fef3c7"
RISK_M_BG   = "fef9c3"
RISK_L_BG   = "dcfce7"

REGULATION_FULL = {
    "GDPR":                 "EU General Data Protection Regulation 2016/679",
    "UK GDPR":              "United Kingdom General Data Protection Regulation",
    "Zimbabwe CDPA":        "Zimbabwe Cyber and Data Protection Act [Chapter 12:07] (2021)",
    "South Africa POPIA":   "Protection of Personal Information Act 4 of 2013 (POPIA)",
    "POPIA":                "Protection of Personal Information Act 4 of 2013 (POPIA)",
    "Kenya DPA":            "Kenya Data Protection Act No. 24 of 2019",
    "Nigeria NDPR":         "Nigeria Data Protection Regulation 2019 / NDPA 2023",
    "Ghana DPA":            "Ghana Data Protection Act 2012 (Act 843)",
    "UAE PDPL":             "UAE Federal Decree-Law No. 45 of 2021 on Personal Data Protection",
    "Saudi PDPL":           "Saudi Arabia Personal Data Protection Law (Royal Decree M/19)",
    "Qatar DPL":            "Qatar Personal Data Privacy Protection Law No. 13 of 2016",
    "Bahrain PDPL":         "Bahrain Personal Data Protection Law 2018",
    "Canada PIPEDA":        "Canada Personal Information Protection and Electronic Documents Act",
    "Canada Bill C-11":     "Canada Consumer Privacy Protection Act (Bill C-11)",
    "USA CCPA/CPRA":        "California Consumer Privacy Act / California Privacy Rights Act",
    "Brazil LGPD":          "Brazil Lei Geral de Proteção de Dados Pessoais — Law No. 13,709/2018",
    "India DPDP":           "India Digital Personal Data Protection Act 2023",
    "Singapore PDPA":       "Singapore Personal Data Protection Act 2012 (revised 2021)",
    "Australia Privacy Act":"Australia Privacy Act 1988 (Cth)",
    "New Zealand Privacy Act":"New Zealand Privacy Act 2020",
    "Japan APPI":           "Japan Act on Protection of Personal Information (APPI) 2022",
    "South Korea PIPA":     "South Korea Personal Information Protection Act (PIPA) 2023",
}

RISK_COLORS = {
    "Very High": (RISK_VH_BG, RISK_VH),
    "High":      (RISK_H_BG,  RISK_H),
    "Medium":    (RISK_M_BG,  RISK_M),
    "Low":       (RISK_L_BG,  RISK_L),
}


def _shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_row_height(row, twips):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(twips))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def _heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.color.rgb = color or (DARK_GREEN if level == 1 else MID_GREEN)
        run.font.name = "Calibri"
    return h


def _info_table(doc, rows, label_width=2.2, value_width=4.7):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    lw = int(label_width * 1440)
    vw = int(value_width * 1440)
    for i, (label, value) in enumerate(rows):
        bg = ALT_ROW if i % 2 else "FFFFFF"
        lc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        lc.width = lw
        vc.width = vw
        _shading(lc, ACCENT_BG)
        _shading(vc, bg)
        lp = lc.paragraphs[0]
        lr = lp.add_run(str(label))
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = DARK_GREEN
        lr.font.name = "Calibri"
        vp = vc.paragraphs[0]
        val_str = str(value) if value else "—"
        vr = vp.add_run(val_str)
        vr.font.size = Pt(9.5)
        vr.font.name = "Calibri"
        vr.font.color.rgb = TEXT_DARK
    doc.add_paragraph()
    return tbl


def _section_header(doc, title, level=1):
    """Adds a shaded section header bar."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _shading(cell, HDR_BG if level == 1 else HDR_BG2)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"
    doc.add_paragraph()


def _risk_badge_table(doc, overall_risk, residual_risk):
    """2-column risk level display."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    labels = [("Overall Risk Level", overall_risk), ("Residual Risk Level", residual_risk)]
    for i, (label, risk) in enumerate(labels):
        cell = tbl.rows[0].cells[i]
        bg, color = RISK_COLORS.get(risk, ("FFFFFF", TEXT_GREY))
        _shading(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}\n")
        r1.font.size = Pt(9)
        r1.font.color.rgb = TEXT_GREY
        r1.font.name = "Calibri"
        r2 = p.add_run(risk or "Not Set")
        r2.bold = True
        r2.font.size = Pt(14)
        r2.font.color.rgb = color
        r2.font.name = "Calibri"
    doc.add_paragraph()


def _render_ai_text(doc, text):
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            for r in p.runs:
                r.font.color.rgb = LIGHT_GREEN
                r.font.name = "Calibri"
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for r in p.runs:
                r.font.color.rgb = MID_GREEN
                r.font.name = "Calibri"
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            for r in p.runs:
                r.font.color.rgb = DARK_GREEN
                r.font.name = "Calibri"
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.paragraph_format.space_before = Pt(2)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.+?)\*\*', stripped)
            for k, part in enumerate(parts):
                if not part:
                    continue
                r = p.add_run(part)
                r.bold = (k % 2 == 1)
                r.font.size = Pt(10.5)
                r.font.name = "Calibri"
                if r.bold:
                    r.font.color.rgb = DARK_GREEN


def generate_docx(dpia):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    regulation      = dpia.get("regulation", "GDPR")
    regulation_full = REGULATION_FULL.get(regulation, regulation)
    created         = dpia.get("created_at", datetime.utcnow().strftime("%Y-%m-%d"))
    ref_num         = dpia.get("ref_number", "")
    title           = dpia.get("title", "Untitled DPIA")

    # ── Cover Page ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    # Shield icon representation using a colored block
    brand_bg = doc.add_table(rows=1, cols=1)
    brand_bg.style = "Table Grid"
    brand_cell = brand_bg.rows[0].cells[0]
    _shading(brand_cell, HDR_BG)
    brand_p = brand_cell.paragraphs[0]
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_p.paragraph_format.space_before = Pt(14)
    brand_p.paragraph_format.space_after  = Pt(14)
    br1 = brand_p.add_run("Data Protection ")
    br1.font.size = Pt(28); br1.font.bold = True; br1.font.color.rgb = WHITE; br1.font.name = "Calibri"
    br2 = brand_p.add_run("Sentinel")
    br2.font.size = Pt(28); br2.font.bold = True; br2.font.color.rgb = RGBColor(0x86, 0xEF, 0xAC); br2.font.name = "Calibri"
    sub_row = brand_cell.add_paragraph()
    sub_row.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_row.add_run("by Ali Moyo  ·  Data Protection Impact Assessment")
    sub_r.font.size = Pt(10); sub_r.font.color.rgb = RGBColor(0xBB, 0xF7, 0xD0); sub_r.font.italic = True
    doc.add_paragraph()

    # Title block
    title_tbl = doc.add_table(rows=1, cols=1)
    title_tbl.style = "Table Grid"
    title_cell = title_tbl.rows[0].cells[0]
    _shading(title_cell, ACCENT_BG)
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after  = Pt(12)
    tr = title_p.add_run(title)
    tr.font.size = Pt(20); tr.font.bold = True; tr.font.color.rgb = DARK_GREEN; tr.font.name = "Calibri"
    reg_p = title_cell.add_paragraph()
    reg_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = reg_p.add_run(regulation_full)
    rr.font.size = Pt(11); rr.font.italic = True; rr.font.color.rgb = TEXT_GREY
    doc.add_paragraph()

    # Cover info table
    overall_risk  = dpia.get("overall_risk", "")
    residual_risk = dpia.get("residual_risk", "")
    overall_bg, overall_col = RISK_COLORS.get(overall_risk, ("FFFFFF", TEXT_GREY))

    _info_table(doc, [
        ("Reference Number",   ref_num),
        ("Organisation",       dpia.get("org_name")),
        ("Department",         dpia.get("department")),
        ("Data Controller",    dpia.get("controller_name")),
        ("DPO / Privacy Lead", dpia.get("dpo_name")),
        ("DPO Email",          dpia.get("dpo_email")),
        ("Activity Type",      dpia.get("activity_type")),
        ("Status",             (dpia.get("status", "draft")).replace("_", " ").title()),
        ("Overall Risk",       overall_risk),
        ("Residual Risk",      residual_risk),
        ("DPO Consulted",      dpia.get("dpo_consulted", "—")),
        ("Authority Consulted",dpia.get("auth_consulted", "—")),
        ("Date Created",       created[:10] if created else "—"),
        ("Regulation",         regulation),
    ])

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = conf.add_run("CONFIDENTIAL — For authorised personnel only")
    cr.font.size = Pt(9); cr.font.color.rgb = TEXT_GREY; cr.font.italic = True

    doc.add_page_break()

    # ── Section 1: Processing Activity ────────────────────────────────────────
    _section_header(doc, "1. Processing Activity Details")
    cats  = dpia.get("data_categories") or "Not specified"
    if isinstance(cats, list):
        cats = ", ".join(cats)
    scats = dpia.get("special_cats") or "None"
    if isinstance(scats, list):
        scats = ", ".join(scats) if scats else "None"
    _info_table(doc, [
        ("Activity Type",      dpia.get("activity_type")),
        ("Activity Description", dpia.get("activity_desc")),
        ("Purpose(s)",         dpia.get("purpose")),
        ("Legal Basis",        dpia.get("legal_basis")),
        ("Data Categories",    cats),
        ("Special Categories", scats),
        ("Data Subjects",      dpia.get("data_subjects")),
        ("Estimated Volume",   dpia.get("subject_count")),
        ("Retention Period",   dpia.get("retention")),
    ])

    # ── Section 2: Systems & Transfers ────────────────────────────────────────
    _section_header(doc, "2. Systems, Processors & International Transfers")
    _info_table(doc, [
        ("Processing Systems",     dpia.get("systems")),
        ("Third-party Processors", dpia.get("processors")),
        ("International Transfer", dpia.get("intl_transfer")),
        ("Destination Countries",  dpia.get("transfer_dest")),
        ("Transfer Mechanism",     dpia.get("transfer_mech")),
    ])

    # ── Section 3: Necessity & Proportionality ────────────────────────────────
    _section_header(doc, "3. Necessity & Proportionality Assessment")

    _heading(doc, "3.1 Necessity Assessment", level=2)
    necessity_text = dpia.get("necessity") or "To be completed."
    doc.add_paragraph(necessity_text)
    doc.add_paragraph()

    _heading(doc, "3.2 Proportionality Assessment", level=2)
    prop_text = dpia.get("proportionality") or "To be completed."
    doc.add_paragraph(prop_text)
    doc.add_paragraph()

    doc.add_page_break()

    # ── Section 4: Risk Assessment ────────────────────────────────────────────
    _section_header(doc, "4. Risk Assessment")
    _risk_badge_table(doc, overall_risk, residual_risk)

    _heading(doc, "4.1 Risk Description & Mitigation", level=2)
    risk_text = dpia.get("risks") or ""
    if isinstance(risk_text, list):
        # Handle old JSON array format
        for r in risk_text:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{r.get('desc', '')} — Likelihood: {r.get('likelihood','')}, Impact: {r.get('impact','')}. Mitigation: {r.get('mitigation','')}")
    elif risk_text:
        doc.add_paragraph(risk_text)
    else:
        doc.add_paragraph("No risks recorded.")
    doc.add_paragraph()

    # ── Section 5: Consultation Record ────────────────────────────────────────
    _section_header(doc, "5. Consultation Record")
    _info_table(doc, [
        ("DPO Consulted",             dpia.get("dpo_consulted") or "No"),
        ("Supervisory Authority",     dpia.get("auth_consulted") or "No"),
        ("Data Subjects Consulted",   dpia.get("subjects_consulted") or "No"),
        ("Consultation Notes",        dpia.get("consult_notes")),
    ])

    # ── Section 6: AI Research (if present) ──────────────────────────────────
    if dpia.get("ai_research"):
        doc.add_page_break()
        _section_header(doc, "6. AI Research Summary")
        note = doc.add_paragraph()
        nr = note.add_run("⚠ AI-assisted research — review and validate before submission.")
        nr.font.size = Pt(9); nr.font.italic = True; nr.font.color.rgb = TEXT_GREY
        doc.add_paragraph()
        _render_ai_text(doc, dpia["ai_research"])

    # ── Section 7: Full AI DPIA Analysis ─────────────────────────────────────
    if dpia.get("ai_full_dpia"):
        doc.add_page_break()
        _section_header(doc, "7. Full DPIA Analysis (AI-Generated)")
        note = doc.add_paragraph()
        nr = note.add_run("⚠ AI-generated content — DPO review required before regulatory submission.")
        nr.font.size = Pt(9); nr.font.italic = True; nr.font.color.rgb = TEXT_GREY
        doc.add_paragraph()
        _render_ai_text(doc, dpia["ai_full_dpia"])

    # ── Sign-Off Page ─────────────────────────────────────────────────────────
    doc.add_page_break()
    _section_header(doc, "Sign-Off & Approval")

    _info_table(doc, [
        ("DPO / Privacy Officer", ""),
        ("Name",       dpia.get("dpo_name") or "________________________________"),
        ("Signature",  "________________________________"),
        ("Date",       "________________________________"),
        ("", ""),
        ("Senior Management", ""),
        ("Name",       "________________________________"),
        ("Title",      "________________________________"),
        ("Signature",  "________________________________"),
        ("Date",       "________________________________"),
    ])

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Generated by Data Protection Sentinel · by Ali Moyo · {datetime.utcnow().strftime('%d %B %Y')}"
    )
    fr.font.size = Pt(8.5); fr.font.color.rgb = TEXT_GREY; fr.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
