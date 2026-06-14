"""
GRID module — Report generation service.

Produces professional branded audit reports in PDF (ReportLab) and DOCX
(python-docx) formats.  Ported from the original AuditSphere pdfReport.js
and wordReport.js with equivalent visual quality.

Security: all inputs are sanitised before embedding in documents.
"""
from __future__ import annotations

import io
import os
import re
import html as html_mod
from datetime import datetime
from core.timeutils import utcnow
from pathlib import Path

# ── PDF (ReportLab) ─────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, HRFlowable,
)
from reportlab.graphics.shapes import Drawing, Rect, String
from reportlab.graphics import renderPDF

# ── DOCX (python-docx) ─────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Local ───────────────────────────────────────────────────────────────
from database import get_db

# ═════════════════════════════════════════════════════════════════════════
# Shared helpers
# ═════════════════════════════════════════════════════════════════════════

# Colour palette (matching original G.R.I.D branding)
DARK   = "#0d0f14"
ACCENT = "#4f8ef7"
GREEN  = "#16a34a"
AMBER  = "#d97706"
RED    = "#dc2626"
MUTED  = "#6b7280"
LIGHT  = "#f3f4f6"

# Reports directory (inside data/)
REPORTS_DIR = Path(os.getenv("REPORTS_DIR", "data/reports"))
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _sanitise(text: str | None) -> str:
    """Strip any potential injection from text destined for documents."""
    if not text:
        return ""
    # Remove HTML tags, control characters
    text = re.sub(r"<[^>]+>", "", str(text))
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text.strip()


def _xml_safe(text: str | None) -> str:
    """Sanitise + escape XML entities for ReportLab Paragraph markup."""
    t = _sanitise(text)
    t = t.replace("&", "&amp;")
    t = t.replace("<", "&lt;")
    t = t.replace(">", "&gt;")
    return t


def _risk_colour(risk: str | None):
    """Return hex colour for risk level."""
    return {
        "Critical": RED, "High": AMBER, "Medium": ACCENT, "Low": GREEN,
    }.get(risk or "", MUTED)


def _status_label(status: str | None, due_date: str | None = None) -> str:
    today = utcnow().strftime("%Y-%m-%d")
    if status == "Complete":
        return "Complete"
    if due_date and due_date < today and status != "Complete":
        return "Overdue"
    if status == "In Progress":
        return "In Progress"
    return "Not Started"


def _status_colour(status: str | None) -> str:
    return {"Complete": GREEN, "In Progress": AMBER}.get(status or "", MUTED)


def _gather_report_data(audit_id: int) -> dict | None:
    """Gather all data needed for a report from the database."""
    db = get_db()
    try:
        audit = db.execute(
            "SELECT a.*, f.name AS framework_name, f.color AS framework_color "
            "FROM grid_audits a LEFT JOIN grid_frameworks f ON a.framework_id=f.id "
            "WHERE a.id=%s", (audit_id,)
        ).fetchone()
        if not audit:
            return None
        audit = dict(audit)

        controls = [dict(r) for r in db.execute("""
            SELECT c.*,
                   u.full_name AS assignee_name,
                   (SELECT COUNT(*) FROM grid_evidence_items WHERE control_id=c.id) AS evidence_required,
                   (SELECT COUNT(*) FROM grid_evidence_files WHERE control_id=c.id) AS evidence_count
            FROM grid_controls c
            LEFT JOIN users u ON c.assignee_id=u.id
            WHERE c.audit_id=%s
            ORDER BY c.control_id
        """, (audit_id,)).fetchall()]

        total = len(controls)
        complete = sum(1 for c in controls if c.get("status") == "Complete")
        overdue_count = sum(
            1 for c in controls
            if c.get("status") != "Complete" and c.get("due_date")
            and c["due_date"] < utcnow().strftime("%Y-%m-%d")
        )
        pending = total - complete
        pct = round(complete / total * 100) if total else 0

        # Critical gaps for AI prompt
        critical_gaps = [
            _sanitise(c.get("name", ""))
            for c in controls
            if c.get("risk_level") in ("Critical", "High")
            and c.get("status") != "Complete"
        ][:5]

        # ── Non-conformances ──────────────────────────────────────────
        ncs = [dict(r) for r in db.execute("""
            SELECT nc.*, u.full_name AS assigned_name,
                   c.control_id AS ctrl_ref, c.name AS control_name
            FROM grid_non_conformances nc
            LEFT JOIN users u ON nc.assigned_to=u.id
            LEFT JOIN grid_controls c ON nc.control_id=c.id
            WHERE nc.audit_id=%s
            ORDER BY CASE nc.severity WHEN 'critical' THEN 1 WHEN 'major' THEN 2
                     WHEN 'minor' THEN 3 ELSE 4 END
        """, (audit_id,)).fetchall()]

        nc_total = len(ncs)
        nc_open = sum(1 for n in ncs if n.get("status") != "closed" and n.get("cap_status") != "Closed")
        nc_closed = nc_total - nc_open

        # ── Evidence summary ─────────────────────────────────────────
        ev_summary = dict(db.execute("""
            SELECT COUNT(*) AS total_files,
                   SUM(CASE WHEN ef.status='Approved' THEN 1 ELSE 0 END) AS approved,
                   SUM(CASE WHEN ef.status='Rejected' THEN 1 ELSE 0 END) AS rejected,
                   SUM(CASE WHEN ef.status='Pending' OR ef.status IS NULL THEN 1 ELSE 0 END) AS pending_review
            FROM grid_evidence_files ef
            JOIN grid_controls c ON ef.control_id=c.id
            WHERE c.audit_id=%s
        """, (audit_id,)).fetchone())

        # ── Management responses ─────────────────────────────────────
        mgmt_responses = [dict(r) for r in db.execute("""
            SELECT nc.title, nc.severity, nc.mgmt_response_status,
                   nc.mgmt_response, nc.response_deadline,
                   u.full_name AS responder_name, nc.mgmt_response_at
            FROM grid_non_conformances nc
            LEFT JOIN users u ON nc.mgmt_response_by=u.id
            WHERE nc.audit_id=%s AND nc.mgmt_response_status IS NOT NULL
            ORDER BY nc.mgmt_response_at DESC
        """, (audit_id,)).fetchall()]

        # ── Sign-offs ────────────────────────────────────────────────
        signoffs = [dict(r) for r in db.execute("""
            SELECT s.role, s.signed_at, s.comment, u.full_name AS user_name
            FROM grid_audit_signoffs s
            LEFT JOIN users u ON s.user_id=u.id
            WHERE s.audit_id=%s
            ORDER BY s.signed_at
        """, (audit_id,)).fetchall()]

        return {
            "audit": audit,
            "controls": controls,
            "auditName": _sanitise(audit.get("name", "Audit Report")),
            "framework": _sanitise(audit.get("framework_name", "")),
            "audit_type": _sanitise(audit.get("audit_type", "External")),
            "auditor": _sanitise(audit.get("auditor", "N/A")),
            "start_date": audit.get("start_date", ""),
            "audit_date": audit.get("audit_date", "TBD"),
            "conclusion": _sanitise(audit.get("conclusion", "")),
            "totalControls": total,
            "complete": complete,
            "pending": pending,
            "overdue": overdue_count,
            "completionPct": pct,
            "criticalGaps": critical_gaps,
            "ncs": ncs,
            "nc_total": nc_total,
            "nc_open": nc_open,
            "nc_closed": nc_closed,
            "evidence_summary": ev_summary,
            "mgmt_responses": mgmt_responses,
            "signoffs": signoffs,
        }
    finally:
        db.close()


# ═════════════════════════════════════════════════════════════════════════
# PDF REPORT GENERATION
# ═════════════════════════════════════════════════════════════════════════

def _hex_to_rl(hex_color: str):
    """Convert #rrggbb to reportlab Color."""
    h = hex_color.lstrip("#")
    return colors.HexColor(f"#{h}")


def generate_pdf_report(audit_id: int, narrative: dict | None = None) -> dict:
    """Generate a branded PDF audit report.

    Returns: {"filePath": str, "fileName": str, "size": int}
    """
    data = _gather_report_data(audit_id)
    if not data:
        raise ValueError(f"Audit {audit_id} not found")

    nar = narrative or {}
    timestamp = int(utcnow().timestamp())
    filename = f"grid-report-{audit_id}-{timestamp}.pdf"
    filepath = REPORTS_DIR / filename

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=50, rightMargin=50, topMargin=60, bottomMargin=50,
        title=f"{data['auditName']} — Audit Report",
        author="G.R.I.D AI — One For All",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Brand", fontName="Helvetica-Bold", fontSize=20, textColor=_hex_to_rl(ACCENT)))
    styles.add(ParagraphStyle("SubBrand", fontName="Helvetica", fontSize=10, textColor=_hex_to_rl(MUTED)))
    styles.add(ParagraphStyle("ReportTitle", fontName="Helvetica-Bold", fontSize=16, textColor=_hex_to_rl(DARK), spaceAfter=6))
    styles.add(ParagraphStyle("SectionHead", fontName="Helvetica-Bold", fontSize=13, textColor=_hex_to_rl(DARK), spaceBefore=16, spaceAfter=8))
    styles.add(ParagraphStyle("Body", fontName="Helvetica", fontSize=10, textColor=_hex_to_rl("#333333"), leading=15, spaceAfter=8))
    styles.add(ParagraphStyle("SmallMuted", fontName="Helvetica", fontSize=9, textColor=_hex_to_rl(MUTED)))
    styles.add(ParagraphStyle("Bullet", fontName="Helvetica", fontSize=10, textColor=_hex_to_rl("#333333"), leading=14, bulletIndent=10, leftIndent=20, spaceAfter=4))
    styles.add(ParagraphStyle("CenterBold", fontName="Helvetica-Bold", fontSize=12, alignment=TA_CENTER, textColor=_hex_to_rl(DARK)))

    elements = []

    # ── Header bar ──────────────────────────────────────────────────
    elements.append(Paragraph("G.R.I.D AI", styles["Brand"]))
    elements.append(Paragraph("Governance · Risk · IT · Data — Compliance Audit Report", styles["SubBrand"]))
    elements.append(Spacer(1, 4))
    elements.append(HRFlowable(width="100%", thickness=1, color=_hex_to_rl(ACCENT)))
    elements.append(Spacer(1, 12))

    # ── Title + metadata ────────────────────────────────────────────
    elements.append(Paragraph(_xml_safe(data["auditName"]), styles["ReportTitle"]))
    meta = (
        f"Framework: {_xml_safe(data['framework'])} · Type: {_xml_safe(data['audit_type'])} · "
        f"Auditor: {_xml_safe(data['auditor'])} · Generated: {utcnow().strftime('%d %b %Y')}"
    )
    elements.append(Paragraph(meta, styles["SmallMuted"]))
    elements.append(Spacer(1, 6))

    # Status pill
    overall = _xml_safe(nar.get("overall_status", "In Progress"))
    st_color = GREEN if overall == "On Track" else AMBER if overall == "At Risk" else RED
    elements.append(Paragraph(
        f'<font color="{st_color}" size="9"><b>{overall.upper()}</b></font>',
        styles["Body"],
    ))
    elements.append(Spacer(1, 12))

    # ── Stats row ───────────────────────────────────────────────────
    stat_data = [
        ["Total Controls", str(data["totalControls"]), ACCENT],
        ["Complete", str(data["complete"]), GREEN],
        ["Pending", str(data["pending"]), AMBER],
        ["Overdue", str(data["overdue"]), RED],
        ["Completion", f"{data['completionPct']}%", ACCENT],
    ]
    stat_table_data = [
        [Paragraph(f'<font color="{c}" size="16"><b>{v}</b></font><br/>'
                   f'<font color="{MUTED}" size="7">{l}</font>', styles["Body"])
         for l, v, c in stat_data]
    ]
    stat_tbl = Table(stat_table_data, colWidths=[doc.width / 5] * 5)
    stat_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _hex_to_rl(LIGHT)),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("BOX", (0, 0), (-1, -1), 0.5, _hex_to_rl("#e5e7eb")),
    ]))
    elements.append(stat_tbl)
    elements.append(Spacer(1, 16))

    # ── Executive Summary ───────────────────────────────────────────
    exec_sum = _xml_safe(nar.get("executive_summary", ""))
    if exec_sum:
        elements.append(Paragraph("Executive Summary", styles["SectionHead"]))
        elements.append(Paragraph(exec_sum, styles["Body"]))
        elements.append(Spacer(1, 8))

    # ── Key Findings ────────────────────────────────────────────────
    findings = nar.get("key_findings", [])
    if findings:
        elements.append(Paragraph("Key Findings", styles["SectionHead"]))
        for f in findings:
            elements.append(Paragraph(f"• {_xml_safe(f)}", styles["Bullet"]))
        elements.append(Spacer(1, 8))

    # ── Controls Table ──────────────────────────────────────────────
    elements.append(Paragraph("Controls Summary", styles["SectionHead"]))
    header = ["Control ID", "Name", "Risk", "Due Date", "Evidence", "Status"]
    tbl_data = [header]

    for c in data["controls"][:60]:  # cap at 60 for PDF
        cid = _xml_safe(c.get("control_id", "—"))
        name = _xml_safe(c.get("name", ""))[:50]
        risk = _xml_safe(c.get("risk_level", "—"))
        due = c.get("due_date", "—") or "—"
        ev = f"{c.get('evidence_count', 0)}/{c.get('evidence_required', 0) or 1}"
        sl = _status_label(c.get("status"), c.get("due_date"))
        tbl_data.append([cid, name, risk, due, ev, sl])

    col_widths = [70, 180, 55, 70, 50, 70]
    tbl = Table(tbl_data, colWidths=col_widths, repeatRows=1)
    tbl_style = [
        # Header
        ("BACKGROUND", (0, 0), (-1, 0), _hex_to_rl(DARK)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, -1), 7),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.3, _hex_to_rl("#e5e7eb")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    # Alternating rows
    for i in range(1, len(tbl_data)):
        if i % 2 == 0:
            tbl_style.append(("BACKGROUND", (0, i), (-1, i), _hex_to_rl("#f8f9fa")))
        # Risk colour
        risk_val = tbl_data[i][2]
        rc = _risk_colour(risk_val)
        tbl_style.append(("TEXTCOLOR", (2, i), (2, i), _hex_to_rl(rc)))
        # Status colour
        sc = _status_colour(tbl_data[i][5].replace("Overdue", "").strip() or tbl_data[i][5])
        if "Overdue" in tbl_data[i][5]:
            sc = RED
        tbl_style.append(("TEXTCOLOR", (5, i), (5, i), _hex_to_rl(sc)))

    tbl.setStyle(TableStyle(tbl_style))
    elements.append(tbl)
    elements.append(Spacer(1, 16))

    # ── Evidence Summary ───────────────────────────────────────────
    ev = data.get("evidence_summary", {})
    if ev.get("total_files"):
        elements.append(Paragraph("Evidence Summary", styles["SectionHead"]))
        ev_data = [
            ["Total Files", "Approved", "Rejected", "Pending Review"],
            [str(ev.get("total_files", 0)), str(ev.get("approved", 0)),
             str(ev.get("rejected", 0)), str(ev.get("pending_review", 0))],
        ]
        ev_tbl = Table(ev_data, colWidths=[doc.width / 4] * 4)
        ev_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), _hex_to_rl(DARK)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("GRID", (0, 0), (-1, -1), 0.3, _hex_to_rl("#e5e7eb")),
        ]))
        elements.append(ev_tbl)
        elements.append(Spacer(1, 16))

    # ── Non-Conformances ───────────────────────────────────────────
    ncs = data.get("ncs", [])
    if ncs:
        elements.append(Paragraph(
            f"Non-Conformances ({data['nc_total']} total — "
            f"{data['nc_open']} open, {data['nc_closed']} closed)",
            styles["SectionHead"],
        ))
        nc_header = ["Title", "Severity", "CAP Status", "Assigned To", "Control"]
        nc_tbl_data = [nc_header]
        for nc in ncs[:40]:
            nc_tbl_data.append([
                _xml_safe(nc.get("title", ""))[:45],
                _xml_safe(nc.get("severity", "—")),
                _xml_safe(nc.get("cap_status", "Open")),
                _xml_safe(nc.get("assigned_name", "—")),
                _xml_safe(nc.get("ctrl_ref", "—")),
            ])
        nc_col_w = [160, 60, 80, 100, 60]
        nc_tbl = Table(nc_tbl_data, colWidths=nc_col_w, repeatRows=1)
        nc_style = [
            ("BACKGROUND", (0, 0), (-1, 0), _hex_to_rl(RED)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("GRID", (0, 0), (-1, -1), 0.3, _hex_to_rl("#e5e7eb")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        for i in range(1, len(nc_tbl_data)):
            if i % 2 == 0:
                nc_style.append(("BACKGROUND", (0, i), (-1, i), _hex_to_rl("#fef2f2")))
            sev = nc_tbl_data[i][1]
            sev_c = RED if sev.lower() == "critical" else AMBER if sev.lower() == "major" else MUTED
            nc_style.append(("TEXTCOLOR", (1, i), (1, i), _hex_to_rl(sev_c)))
        nc_tbl.setStyle(TableStyle(nc_style))
        elements.append(nc_tbl)
        elements.append(Spacer(1, 16))

    # ── Management Responses ───────────────────────────────────────
    mgmt = data.get("mgmt_responses", [])
    if mgmt:
        elements.append(Paragraph("Management Responses", styles["SectionHead"]))
        for m in mgmt[:10]:
            status = _xml_safe(m.get("mgmt_response_status", ""))
            sc = GREEN if status == "Approved" else RED
            elements.append(Paragraph(
                f'<font color="{sc}"><b>{status}</b></font> — '
                f'{_xml_safe(m.get("title", ""))} '
                f'<font color="{MUTED}" size="8">(by {_xml_safe(m.get("responder_name", "—"))})</font>',
                styles["Body"],
            ))
            resp_text = _xml_safe(m.get("mgmt_response", ""))
            if resp_text:
                elements.append(Paragraph(
                    f'<font color="{MUTED}" size="8"><i>"{resp_text[:200]}"</i></font>',
                    styles["Body"],
                ))

    # ── Conclusion ──────────────────────────────────────────────────
    # Use audit conclusion first, fall back to AI-generated
    conclusion = _xml_safe(data.get("conclusion", "")) or _xml_safe(nar.get("conclusion", ""))
    if conclusion:
        elements.append(Paragraph("Conclusion", styles["SectionHead"]))
        elements.append(Paragraph(conclusion, styles["Body"]))

    # ── Sign-offs ──────────────────────────────────────────────────
    signoffs = data.get("signoffs", [])
    if signoffs:
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Sign-offs", styles["SectionHead"]))
        for s in signoffs:
            elements.append(Paragraph(
                f'<b>{_xml_safe(s.get("role", "").title())}</b>: '
                f'{_xml_safe(s.get("user_name", "—"))} — '
                f'{_xml_safe(s.get("signed_at", "")[:10])}'
                + (f'  <font color="{MUTED}" size="8"><i>"{_xml_safe(s.get("comment", ""))}"</i></font>'
                   if s.get("comment") else ""),
                styles["Body"],
            ))

    # ── Footer line ─────────────────────────────────────────────────
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=_hex_to_rl("#e5e7eb")))
    elements.append(Paragraph(
        f"G.R.I.D AI — One For All Compliance Platform · Confidential · "
        f"{utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        styles["SmallMuted"],
    ))

    doc.build(elements)
    pdf_bytes = buf.getvalue()
    filepath.write_bytes(pdf_bytes)

    return {"filePath": str(filepath), "fileName": filename, "size": len(pdf_bytes)}


# ═════════════════════════════════════════════════════════════════════════
# DOCX REPORT GENERATION
# ═════════════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a docx table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_run(paragraph, text: str, bold=False, size=None, color=None, italic=False):
    """Add a run to a paragraph with formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        h = color.lstrip("#")
        run.font.color.rgb = RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return run


def generate_docx_report(audit_id: int, narrative: dict | None = None) -> dict:
    """Generate a branded DOCX audit report.

    Returns: {"filePath": str, "fileName": str, "size": int}
    """
    data = _gather_report_data(audit_id)
    if not data:
        raise ValueError(f"Audit {audit_id} not found")

    nar = narrative or {}
    timestamp = int(utcnow().timestamp())
    filename = f"grid-report-{audit_id}-{timestamp}.docx"
    filepath = REPORTS_DIR / filename
    gen_date = utcnow().strftime("%d %B %Y")

    doc = DocxDocument()

    # Page setup — A4 landscape for better table layout
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # ── COVER PAGE ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(80)
    _add_run(p, "G.R.I.D AI", bold=True, size=28, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Governance · Risk · IT · Data", size=12, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(6)
    _add_run(p, "by One For All Compliance Platform", size=9, color=MUTED)

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    _add_run(p, "AUDIT REPORT", bold=True, size=24, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    _add_run(p, _sanitise(data["auditName"]), bold=True, size=16, color="#374151")

    # Metadata table
    doc.add_paragraph()
    meta_items = [
        ("Framework:", data["framework"]),
        ("Audit Type:", data["audit_type"]),
        ("Auditor:", data["auditor"]),
        ("Audit Date:", data["audit_date"] or "TBD"),
        ("Generated:", gen_date),
        ("Status:", _sanitise(nar.get("overall_status", "In Progress"))),
    ]
    meta_table = doc.add_table(rows=len(meta_items), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_items):
        c0 = meta_table.cell(i, 0)
        c1 = meta_table.cell(i, 1)
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        _add_run(p0, label, bold=True, size=9, color=MUTED)
        color = DARK
        if label == "Status:":
            s = value
            color = GREEN if s == "On Track" else AMBER if s == "At Risk" else RED
        _add_run(p1, value, size=9, color=color)
        # Remove borders
        for cell in (c0, c1):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders_el = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                borders_el.append(el)
            tcPr.append(borders_el)

    doc.add_page_break()

    # ── STATS SUMMARY ───────────────────────────────────────────────
    doc.add_heading("Audit At a Glance", level=1)

    stats_items = [
        ("Total Controls", str(data["totalControls"]), ACCENT),
        ("Complete", str(data["complete"]), GREEN),
        ("Pending", str(data["pending"]), AMBER),
        ("Overdue", str(data["overdue"]), RED),
        ("Completion", f"{data['completionPct']}%", ACCENT),
    ]
    stats_table = doc.add_table(rows=2, cols=5)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, color) in enumerate(stats_items):
        # Value row
        cell_v = stats_table.cell(0, i)
        pv = cell_v.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(pv, value, bold=True, size=18, color=color)
        _set_cell_shading(cell_v, LIGHT)
        # Label row
        cell_l = stats_table.cell(1, i)
        pl = cell_l.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(pl, label, size=8, color=MUTED)
        _set_cell_shading(cell_l, LIGHT)
    # Remove table borders
    for row in stats_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders_el = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                borders_el.append(el)
            tcPr.append(borders_el)

    doc.add_paragraph()

    # ── EXECUTIVE SUMMARY ───────────────────────────────────────────
    exec_sum = _sanitise(nar.get("executive_summary", ""))
    if exec_sum:
        doc.add_heading("Executive Summary", level=1)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        _add_run(p, exec_sum, size=10, color="#374151")
        # Blue left border via XML
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left_b = OxmlElement("w:left")
        left_b.set(qn("w:val"), "single")
        left_b.set(qn("w:sz"), "18")
        left_b.set(qn("w:color"), ACCENT.lstrip("#"))
        left_b.set(qn("w:space"), "8")
        pBdr.append(left_b)
        pPr.append(pBdr)

    # ── KEY FINDINGS ────────────────────────────────────────────────
    findings = nar.get("key_findings", [])
    if findings:
        doc.add_heading("Key Findings", level=2)
        for f in findings:
            p = doc.add_paragraph(style="List Bullet")
            _add_run(p, _sanitise(f), size=10, color="#374151")

    doc.add_page_break()

    # ── CONTROLS TABLE ──────────────────────────────────────────────
    doc.add_heading("Controls Summary", level=1)

    cols = ["Control ID", "Control Name", "Risk", "Due Date", "Evidence", "Status"]
    ctrl_table = doc.add_table(rows=1 + len(data["controls"]), cols=len(cols))
    ctrl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ctrl_table.autofit = True

    # Header row
    for j, col_name in enumerate(cols):
        cell = ctrl_table.cell(0, j)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, col_name, bold=True, size=8, color="#ffffff")
        _set_cell_shading(cell, DARK)

    # Data rows
    for i, c in enumerate(data["controls"]):
        row_idx = i + 1
        bg = "#ffffff" if i % 2 == 0 else LIGHT

        cid = _sanitise(c.get("control_id", "—"))
        name = _sanitise(c.get("name", ""))[:60]
        risk = _sanitise(c.get("risk_level", "—"))
        due = c.get("due_date", "—") or "—"
        ev_req = c.get("evidence_required", 0) or 1
        ev_cnt = c.get("evidence_count", 0)
        sl = _status_label(c.get("status"), c.get("due_date"))

        row_data = [
            (cid, ACCENT, True),
            (name, "#374151", False),
            (risk, _risk_colour(risk), True),
            (due, "#374151", False),
            (f"{ev_cnt}/{ev_req}", GREEN if ev_cnt >= ev_req else AMBER, False),
            (sl, RED if sl == "Overdue" else _status_colour(c.get("status")), True),
        ]

        for j, (text, color, bold) in enumerate(row_data):
            cell = ctrl_table.cell(row_idx, j)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            _add_run(p, text, bold=bold, size=8, color=color)
            _set_cell_shading(cell, bg)

    # Apply thin borders to all cells
    tbl_xml = ctrl_table._tbl
    tblPr = tbl_xml.tblPr if tbl_xml.tblPr is not None else OxmlElement("w:tblPr")
    borders_el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "e5e7eb")
        el.set(qn("w:space"), "0")
        borders_el.append(el)
    tblPr.append(borders_el)

    # ── EVIDENCE SUMMARY ───────────────────────────────────────────
    ev = data.get("evidence_summary", {})
    if ev.get("total_files"):
        doc.add_heading("Evidence Summary", level=2)
        ev_tbl = doc.add_table(rows=2, cols=4)
        ev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, (label, val) in enumerate([
            ("Total Files", str(ev.get("total_files", 0))),
            ("Approved", str(ev.get("approved", 0))),
            ("Rejected", str(ev.get("rejected", 0))),
            ("Pending", str(ev.get("pending_review", 0))),
        ]):
            cell_h = ev_tbl.cell(0, j)
            _add_run(cell_h.paragraphs[0], label, bold=True, size=8, color="#ffffff")
            cell_h.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell_h, DARK)
            cell_v = ev_tbl.cell(1, j)
            _add_run(cell_v.paragraphs[0], val, bold=True, size=12, color=ACCENT)
            cell_v.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ── NON-CONFORMANCES ───────────────────────────────────────────
    ncs = data.get("ncs", [])
    if ncs:
        doc.add_heading(
            f"Non-Conformances ({data['nc_total']} total — "
            f"{data['nc_open']} open, {data['nc_closed']} closed)",
            level=1,
        )
        nc_cols = ["Title", "Severity", "CAP Status", "Assigned To", "Control"]
        nc_tbl = doc.add_table(rows=1 + min(len(ncs), 40), cols=len(nc_cols))
        nc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        nc_tbl.autofit = True
        for j, col_name in enumerate(nc_cols):
            cell = nc_tbl.cell(0, j)
            _add_run(cell.paragraphs[0], col_name, bold=True, size=8, color="#ffffff")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, RED.lstrip("#") if True else DARK)
            _set_cell_shading(cell, RED)
        for i, nc in enumerate(ncs[:40]):
            row_idx = i + 1
            bg = "#ffffff" if i % 2 == 0 else "#fef2f2"
            sev = _sanitise(nc.get("severity", "—"))
            sev_c = RED if sev.lower() == "critical" else AMBER if sev.lower() == "major" else MUTED
            vals = [
                (_sanitise(nc.get("title", ""))[:45], "#374151", False),
                (sev, sev_c, True),
                (_sanitise(nc.get("cap_status", "Open")), ACCENT, True),
                (_sanitise(nc.get("assigned_name", "—")), "#374151", False),
                (_sanitise(nc.get("ctrl_ref", "—")), MUTED, False),
            ]
            for j, (text, color, bold) in enumerate(vals):
                cell = nc_tbl.cell(row_idx, j)
                _add_run(cell.paragraphs[0], text, bold=bold, size=8, color=color)
                _set_cell_shading(cell, bg)
        doc.add_paragraph()

    # ── MANAGEMENT RESPONSES ───────────────────────────────────────
    mgmt = data.get("mgmt_responses", [])
    if mgmt:
        doc.add_heading("Management Responses", level=2)
        for m in mgmt[:10]:
            status = _sanitise(m.get("mgmt_response_status", ""))
            sc = GREEN if status == "Approved" else RED
            p = doc.add_paragraph()
            _add_run(p, f"{status}", bold=True, size=10, color=sc)
            _add_run(p, f" — {_sanitise(m.get('title', ''))}", size=10, color="#374151")
            _add_run(p, f"  (by {_sanitise(m.get('responder_name', '—'))})", size=8, color=MUTED)
            resp_text = _sanitise(m.get("mgmt_response", ""))
            if resp_text:
                pr = doc.add_paragraph()
                _add_run(pr, f'"{resp_text[:200]}"', size=8, color=MUTED, italic=True)

    # ── CONCLUSION ─────────────────────────────────────────────────
    conclusion_text = data.get("conclusion") or _sanitise(nar.get("conclusion", ""))
    if conclusion_text:
        doc.add_heading("Conclusion", level=1)
        p = doc.add_paragraph()
        _add_run(p, conclusion_text, size=10, color="#374151")

    # ── SIGN-OFFS ──────────────────────────────────────────────────
    signoffs = data.get("signoffs", [])
    if signoffs:
        doc.add_heading("Sign-offs", level=2)
        for s in signoffs:
            p = doc.add_paragraph()
            _add_run(p, f"{_sanitise(s.get('role', '').title())}: ", bold=True, size=10, color=DARK)
            _add_run(p, _sanitise(s.get("user_name", "—")), size=10, color="#374151")
            _add_run(p, f" — {_sanitise(s.get('signed_at', '')[:10])}", size=9, color=MUTED)
            if s.get("comment"):
                _add_run(p, f'  "{_sanitise(s["comment"])}"', size=8, color=MUTED, italic=True)

    # ── HEADER & FOOTER ─────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    _add_run(hp, "G.R.I.D AI  ", bold=True, size=9, color=ACCENT)
    _add_run(hp, f"{data['auditName']} — {data['framework']} Audit Report", size=9, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(fp, f"Generated by G.R.I.D AI — One For All  ·  {gen_date}  ·  Confidential", size=8, color=MUTED)

    # ── SAVE ────────────────────────────────────────────────────────
    doc.save(str(filepath))
    size = filepath.stat().st_size

    return {"filePath": str(filepath), "fileName": filename, "size": size}
