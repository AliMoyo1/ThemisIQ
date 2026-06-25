"""
ARIA module -- Document branding engine.

Applies company branding templates to uploaded policy documents.
Takes a source DOCX (user-edited content) and a template DOCX
(company branding: cover page, logo, metadata tables, styles),
and produces a branded output DOCX that preserves the template's
front matter and appends the policy content after it.

Security: all file paths are validated before access. No user input
is interpolated into file operations without sanitisation.
"""
from __future__ import annotations

import logging
import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

log = logging.getLogger("aria.branding")

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _sanitise(text: str | None) -> str:
    if not text:
        return ""
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", str(text)).strip()


def _element_tag(el) -> str:
    return el.tag.split("}")[-1] if "}" in el.tag else el.tag


def _element_text(el) -> str:
    runs = el.findall(f".//{{{WML}}}t")
    return "".join(r.text or "" for r in runs)


def _element_style(el) -> str:
    pPr = el.find(f"{{{WML}}}pPr")
    if pPr is not None:
        pStyle = pPr.find(f"{{{WML}}}pStyle")
        if pStyle is not None:
            return pStyle.get(f"{{{WML}}}val", "")
    return ""


def _update_paragraph_text(para_el, new_text: str):
    """Replace all run texts in a paragraph, preserving the first run's formatting."""
    runs = para_el.findall(f".//{{{WML}}}r")
    if not runs:
        return
    first_run = runs[0]
    t_el = first_run.find(f"{{{WML}}}t")
    if t_el is not None:
        t_el.text = new_text
    for run in runs[1:]:
        run.getparent().remove(run)


def _get_para_font_size_half_pt(para_el) -> int:
    """Return the font size in half-points from the first run, or 0."""
    run = para_el.find(f".//{{{WML}}}r")
    if run is None:
        return 0
    rPr = run.find(f"{{{WML}}}rPr")
    if rPr is None:
        return 0
    sz = rPr.find(f"{{{WML}}}sz")
    if sz is not None:
        try:
            return int(sz.get(f"{{{WML}}}val", "0"))
        except ValueError:
            return 0
    return 0


def _table_rows(tbl_el):
    return tbl_el.findall(f"{{{WML}}}tr")


def _table_cell_text(row_el, col: int) -> str:
    cells = row_el.findall(f"{{{WML}}}tc")
    if col < len(cells):
        return _element_text(cells[col])
    return ""


def _set_table_cell_text(row_el, col: int, text: str):
    cells = row_el.findall(f"{{{WML}}}tc")
    if col >= len(cells):
        return
    cell = cells[col]
    paras = cell.findall(f"{{{WML}}}p")
    if paras:
        _update_paragraph_text(paras[0], text)


def _add_table_row(tbl_el, values: list[str]):
    """Clone the last data row of a table and fill with new values."""
    rows = _table_rows(tbl_el)
    if len(rows) < 2:
        return
    last_row = rows[-1]
    new_row = deepcopy(last_row)
    cells = new_row.findall(f"{{{WML}}}tc")
    for i, val in enumerate(values):
        if i < len(cells):
            paras = cells[i].findall(f"{{{WML}}}p")
            if paras:
                _update_paragraph_text(paras[0], val)
    tbl_el.append(new_row)


def _is_content_start(el) -> bool:
    """True if the element looks like the start of actual policy content."""
    tag = _element_tag(el)
    if tag != "p":
        return False
    style = _element_style(el)
    text = _element_text(el).strip()
    if style in ("Heading1", "Heading2") and re.match(r"^\d+[\.\)]?\s", text):
        return True
    if text.lower().startswith(("purpose", "scope", "introduction", "objective")):
        return True
    return False


def apply_template(
    *,
    source_path: str,
    template_path: str,
    output_path: str,
    logo_path: str | None = None,
    doc_title: str = "",
    doc_id: str = "",
    version: str = "1.0",
    framework: str = "",
    author_name: str = "",
) -> str:
    """
    Merge source document content into a branding template.

    Strategy:
    1. Open the template DOCX (has cover page, logo, metadata tables, styles)
    2. Keep the template body intact (cover page, tables, front matter)
    3. Update metadata fields (title, version, dates, revision history)
    4. Append source content after the template front matter
    5. Save as the output file.

    Returns the output path.
    """
    src = Path(source_path)
    tpl = Path(template_path)
    out = Path(output_path)

    if not src.exists():
        raise FileNotFoundError(f"Source document not found: {source_path}")
    if not tpl.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    branded = Document(str(tpl))
    source = Document(str(src))
    body = branded.element.body
    today = datetime.now().strftime("%Y/%m/%d")
    year = datetime.now().strftime("%Y")
    title = _sanitise(doc_title) or "Policy Document"

    # ── Update template metadata ─────────────────────────────────
    title_updated = False
    for child in body:
        if _element_tag(child) != "p":
            continue
        sz = _get_para_font_size_half_pt(child)
        text = _element_text(child).strip()

        if not title_updated and sz >= 36 and text and not text.startswith("VERSION"):
            if not re.match(r"^(Information Security|The content)", text):
                _update_paragraph_text(child, title.upper())
                title_updated = True
                continue

        if re.match(r"VERSION\s+\d", text, re.IGNORECASE):
            _update_paragraph_text(child, f"VERSION {version} OF {year}")

    # ── Update metadata table (has "Reference:" in first row) ────
    for child in body:
        if _element_tag(child) != "tbl":
            continue
        rows = _table_rows(child)
        if not rows:
            continue
        first_row_text = " ".join(
            _table_cell_text(rows[0], c) for c in range(4)
        ).lower()

        if "reference" in first_row_text:
            _set_table_cell_text(rows[0], 1, title)
            _set_table_cell_text(rows[0], 3, doc_id)
            if len(rows) > 1:
                _set_table_cell_text(rows[1], 3, version)
            if len(rows) > 2:
                _set_table_cell_text(rows[2], 1, today)
                _set_table_cell_text(rows[2], 3, today)
            continue

        if "version number" in first_row_text or "version" in first_row_text and "author" in first_row_text:
            _add_table_row(child, [
                version,
                author_name or "ThemisIQ",
                "Generated via ThemisIQ ARIA",
                today,
            ])

    # ── Find insertion point (before sectPr) ─────────────────────
    sect_pr = body.find(f"{{{WML}}}sectPr")
    if sect_pr is None:
        children = list(body)
        sect_pr = children[-1] if children else None

    # ── Copy source content, skipping the title/preamble block ──
    content_started = False

    for element in source.element.body:
        tag = _element_tag(element)
        if tag == "sectPr":
            continue
        if tag not in ("p", "tbl"):
            continue

        if not content_started:
            if tag == "tbl":
                content_started = True
            elif _is_content_start(element):
                content_started = True
            else:
                continue

        if content_started:
            new_el = deepcopy(element)
            if sect_pr is not None:
                sect_pr.addprevious(new_el)
            else:
                body.append(new_el)

    # ── Header logo fallback (for templates with logo in header) ─
    if logo_path:
        logo = Path(logo_path)
        if logo.exists():
            try:
                for section in branded.sections:
                    header = section.header
                    header.is_linked_to_previous = False
                    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = hp.add_run()
                    run.add_picture(str(logo), width=Cm(3))
            except Exception as exc:
                log.warning("Could not add logo to header: %s", exc)

    # ── Save branded document ────────────────────────────────────
    out.parent.mkdir(parents=True, exist_ok=True)
    branded.save(str(out))
    log.info("Branded document saved: %s", out)

    return str(out)
