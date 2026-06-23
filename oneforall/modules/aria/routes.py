"""
ARIA module -- Governance, Risk & Compliance.

Full routes: dashboard, frameworks, controls, documents, risks,
cross-mapping, export, AI generator, Ask ARIA, audit log.
"""
import io
import json
import logging
import re as _re
from datetime import datetime, timedelta
from core.timeutils import utcnow
from typing import Optional

log = logging.getLogger("oneforall.aria")

from fastapi import APIRouter, Request, Form, HTTPException, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, StreamingResponse
from fastapi.templating import Jinja2Templates

from config import settings
from database import get_db, insert_returning_id
from core.middleware import (
    get_current_user, require_auth, require_module,
    require_capability, log_audit,
)
from core.rbac import (
    has_capability, has_role, ROLE_LABELS, ALL_ROLES,
    CAPABILITIES, user_capabilities,
)

from core.shell_context import shell_ctx
from core.events import (
    emit, ARIA_POLICY_PUBLISHED, ARIA_POLICY_UPDATED,
    ARIA_RISK_CREATED, ARIA_RISK_ESCALATED, ARIA_CONTROL_UPDATED,
)

router = APIRouter(prefix="/aria", tags=["aria"])
templates = Jinja2Templates(directory=["modules/aria/templates", "templates"])

# Expose helpers to all ARIA templates
templates.env.globals["has_capability"] = has_capability
templates.env.globals["has_role"] = has_role
templates.env.globals["ROLE_LABELS"] = ROLE_LABELS
templates.env.filters["tojson"] = lambda v: json.dumps(
    dict(v) if hasattr(v, "keys") else v, default=str
)


def _aria_render(request: Request, template: str, context: dict,
                  active_section: str = ""):
    """Render an ARIA template with unified shell context injected."""
    context.setdefault("user", request.state.user)
    context.update(shell_ctx(request, active_module="aria",
                             active_section=active_section))
    return templates.TemplateResponse(request, template, context)

# -- Cross-framework mapping (static reference data) -------------------------

CROSS_MAPPING = {
    "Access Control": {
        "ISO 27001": ["A.5.15", "A.5.16", "A.5.17", "A.5.18", "A.8.2", "A.8.3", "A.8.5"],
        "SOC 2 Type II": ["CC6.1", "CC6.2", "CC6.3", "CC6.4", "CC6.5"],
        "PCI DSS": ["7.1", "7.2", "8.1", "8.2", "8.3", "8.4"],
        "GDPR": ["Art.25", "Art.32"],
        "Zimbabwe CDPA": ["S.11"],
        "HIPAA": ["164.308(a)(4)", "164.312(a)"],
        "ISO 42001": [],
    },
    "Incident Response": {
        "ISO 27001": ["A.5.24", "A.5.25", "A.5.26", "A.5.27", "A.5.28"],
        "SOC 2 Type II": ["CC7.3", "CC7.4", "CC7.5"],
        "PCI DSS": ["12.10"],
        "GDPR": ["Art.33", "Art.34"],
        "Zimbabwe CDPA": ["S.12"],
        "HIPAA": ["164.308(a)(6)", "164.404", "164.406", "164.408"],
        "ISO 42001": ["A.9.2"],
    },
    "Risk Assessment": {
        "ISO 27001": ["6.1.2", "6.1.3"],
        "SOC 2 Type II": ["CC3.1", "CC3.2", "CC3.3", "CC3.4"],
        "PCI DSS": ["12.3"],
        "GDPR": ["Art.35"],
        "Zimbabwe CDPA": ["S.11"],
        "HIPAA": ["164.308(a)(1)"],
        "ISO 42001": ["6.1", "6.1.2", "A.4.3"],
    },
    "Data Classification": {
        "ISO 27001": ["A.5.12", "A.5.13"],
        "SOC 2 Type II": ["C1.1"],
        "PCI DSS": ["3.1", "3.2"],
        "GDPR": ["Art.5", "Art.9"],
        "Zimbabwe CDPA": ["S.6", "S.16"],
        "HIPAA": ["164.514"],
        "ISO 42001": ["A.6.1", "A.6.2"],
    },
    "Third-Party Management": {
        "ISO 27001": ["A.5.19", "A.5.20", "A.5.21", "A.5.22"],
        "SOC 2 Type II": ["CC9.2"],
        "PCI DSS": ["12.8", "12.9"],
        "GDPR": ["Art.28", "Art.29"],
        "Zimbabwe CDPA": ["S.18", "S.19"],
        "HIPAA": ["164.308(b)"],
        "ISO 42001": ["8.7"],
    },
    "Awareness & Training": {
        "ISO 27001": ["7.2", "7.3", "A.6.3"],
        "SOC 2 Type II": ["CC1.4"],
        "PCI DSS": ["12.6"],
        "GDPR": [],
        "Zimbabwe CDPA": [],
        "HIPAA": ["164.308(a)(5)"],
        "ISO 42001": ["7.2", "7.3", "A.2.4", "A.2.5"],
    },
    "Logging & Monitoring": {
        "ISO 27001": ["A.8.15", "A.8.16"],
        "SOC 2 Type II": ["CC4.1", "CC7.1", "CC7.2"],
        "PCI DSS": ["10.1", "10.2", "10.3", "10.4", "10.5"],
        "GDPR": ["Art.30"],
        "Zimbabwe CDPA": [],
        "HIPAA": ["164.312(b)"],
        "ISO 42001": ["A.8.2"],
    },
    "Encryption": {
        "ISO 27001": ["A.8.24"],
        "SOC 2 Type II": ["CC6.7"],
        "PCI DSS": ["3.4", "3.5", "4.1", "4.2"],
        "GDPR": ["Art.32"],
        "Zimbabwe CDPA": ["S.11"],
        "HIPAA": ["164.312(e)"],
        "ISO 42001": [],
    },
}

FRAMEWORK_NAMES = [
    "ISO 27001", "SOC 2 Type II", "PCI DSS",
    "GDPR", "Zimbabwe CDPA", "HIPAA", "ISO 42001",
]


# -- Helpers ------------------------------------------------------------------

def _can_edit_control(user, control=None):
    """Check if user can edit a control."""
    if has_capability(user, "aria.control.update_any"):
        return True
    if has_capability(user, "aria.control.update_own") and control:
        owner = (control.get("owner") or "").strip().lower()
        uname = (user.get("username") or "").strip().lower()
        fname = (user.get("full_name") or "").strip().lower()
        if owner and owner in (uname, fname):
            return True
    return False


def _can_approve_policy(user, doc):
    """Check approval capability with separation-of-duties enforcement."""
    if not has_capability(user, "aria.policy.approve"):
        return False
    if has_capability(user, "platform.manage_users"):
        return True
    doc_owner = (doc.get("owner") or "").strip().lower()
    uname = (user.get("username") or "").strip().lower()
    fname = (user.get("full_name") or "").strip().lower()
    if doc_owner and doc_owner in (uname, fname):
        return False
    return True


# -- Dashboard ----------------------------------------------------------------

@router.get("/", response_class=HTMLResponse)
@require_module("aria")
async def aria_dashboard(request: Request):
    user = request.state.user
    db = get_db()
    try:
        frameworks = db.execute(
            "SELECT * FROM frameworks WHERE is_active = 1 ORDER BY name"
        ).fetchall()
        fw_stats = []
        for fw in frameworks:
            stats = db.execute("""
                SELECT
                    COUNT(*) as total,
                    SUM(CASE WHEN status='Implemented' THEN 1 ELSE 0 END) as implemented,
                    SUM(CASE WHEN status='In Progress' THEN 1 ELSE 0 END) as in_progress,
                    SUM(CASE WHEN status='Not Started' THEN 1 ELSE 0 END) as not_started,
                    SUM(CASE WHEN status='Approved' THEN 1 ELSE 0 END) as approved,
                    SUM(CASE WHEN status='Draft' THEN 1 ELSE 0 END) as draft,
                    SUM(CASE WHEN status='Under Review' THEN 1 ELSE 0 END) as under_review,
                    SUM(CASE WHEN status='Needs Update' THEN 1 ELSE 0 END) as needs_update
                FROM controls WHERE framework_id=%s
            """, (fw["id"],)).fetchone()
            total = stats["total"] or 0
            impl = stats["implemented"] or 0
            pct = round((impl / total * 100) if total > 0 else 0, 1)
            fw_stats.append({
                "id": fw["id"], "name": fw["name"], "color": fw["color"],
                "description": fw["description"],
                "total": total, "implemented": impl,
                "in_progress": stats["in_progress"] or 0,
                "not_started": stats["not_started"] or 0,
                "approved": stats["approved"] or 0,
                "draft": stats["draft"] or 0,
                "under_review": stats["under_review"] or 0,
                "needs_update": stats["needs_update"] or 0,
                "pct": pct,
            })
        totals = {
            "total": sum(f["total"] for f in fw_stats),
            "implemented": sum(f["implemented"] for f in fw_stats),
            "in_progress": sum(f["in_progress"] for f in fw_stats),
            "not_started": sum(f["not_started"] for f in fw_stats),
            "overall_pct": round(
                sum(f["implemented"] for f in fw_stats)
                / max(sum(f["total"] for f in fw_stats), 1) * 100, 1
            ),
        }
        recent = db.execute("""
            SELECT al.*, u.full_name FROM audit_log al
            LEFT JOIN users u ON al.user_id = u.id
            WHERE al.module = 'aria'
            ORDER BY al.created_at DESC LIMIT 8
        """).fetchall()
    finally:
        db.close()

    return _aria_render(request, "dashboard.html", {
        "user": user,
        "module": "aria",
        "fw_stats": fw_stats,
        "fw_stats_json": json.dumps(fw_stats),
        "totals": totals,
        "totals_json": json.dumps(totals),
        "recent": recent,
    }, active_section="dashboard")


# -- Frameworks List ----------------------------------------------------------

@router.get("/frameworks", response_class=HTMLResponse)
@require_module("aria")
async def frameworks_list(request: Request):
    user = request.state.user
    db = get_db()
    try:
        frameworks = db.execute(
            "SELECT * FROM frameworks WHERE is_active = 1 ORDER BY name"
        ).fetchall()

        # Document coverage per framework
        doc_counts = {}
        try:
            doc_rows = db.execute(
                "SELECT framework, COUNT(*) as cnt, "
                "COUNT(DISTINCT control_ref) as ctrl_covered "
                "FROM aria_documents GROUP BY framework"
            ).fetchall()
            for dr in doc_rows:
                doc_counts[dr["framework"]] = {
                    "docs": dr["cnt"], "ctrls_covered": dr["ctrl_covered"]
                }
        except Exception:
            pass

        # Evidence counts per framework
        ev_counts = {}
        try:
            ev_rows = db.execute(
                "SELECT c.framework_id, COUNT(DISTINCT el.id) as ev_count "
                "FROM evidence_links el "
                "JOIN controls c ON c.id = el.entity_id "
                "WHERE el.module = 'aria' AND el.entity_type = 'control' "
                "GROUP BY c.framework_id"
            ).fetchall()
            for er in ev_rows:
                ev_counts[er["framework_id"]] = er["ev_count"]
        except Exception:
            pass

        # IMS mapping counts per framework (how many controls are mapped to another fw)
        mapped_counts = {}
        try:
            map_rows = db.execute(
                "SELECT c.framework_id, COUNT(DISTINCT c.id) as mapped "
                "FROM controls c "
                "WHERE c.framework_id IN (SELECT id FROM frameworks WHERE is_active=1) "
                "AND (c.id IN (SELECT source_control_id FROM aria_control_mappings) "
                "     OR c.id IN (SELECT target_control_id FROM aria_control_mappings)) "
                "GROUP BY c.framework_id"
            ).fetchall()
            for mr in map_rows:
                mapped_counts[mr["framework_id"]] = mr["mapped"]
        except Exception:
            pass

        fw_list = []
        fw_ids = [fw["id"] for fw in frameworks]
        for fw in frameworks:
            stats = db.execute("""
                SELECT COUNT(*) as total,
                    SUM(CASE WHEN status='Implemented' THEN 1 ELSE 0 END) as implemented,
                    SUM(CASE WHEN status='In Progress' THEN 1 ELSE 0 END) as in_progress,
                    SUM(CASE WHEN status='Not Started' THEN 1 ELSE 0 END) as not_started
                FROM controls WHERE framework_id=%s
            """, (fw["id"],)).fetchone()
            total = stats["total"] or 0
            impl = stats["implemented"] or 0
            pct = round((impl / total * 100) if total > 0 else 0, 1)
            dc = doc_counts.get(fw["name"], {})
            doc_total = dc.get("docs", 0)
            ctrls_with_docs = dc.get("ctrls_covered", 0)
            doc_pct = round((ctrls_with_docs / total * 100) if total > 0 else 0)
            ev_total = ev_counts.get(fw["id"], 0)
            mapped = mapped_counts.get(fw["id"], 0)
            fw_list.append({
                "id": fw["id"], "name": fw["name"], "color": fw["color"],
                "description": fw["description"], "total": total,
                "implemented": impl, "in_progress": stats["in_progress"] or 0,
                "not_started": stats["not_started"] or 0, "pct": pct,
                "doc_total": doc_total, "doc_pct": doc_pct,
                "ctrls_with_docs": ctrls_with_docs,
                "evidence_count": ev_total, "mapped_controls": mapped,
            })

        # IMS summary: cross-framework mapping stats
        ims_summary = None
        if len(fw_ids) >= 2:
            try:
                total_mappings = db.execute(
                    "SELECT COUNT(*) FROM aria_control_mappings"
                ).fetchone()[0]
                total_ctrls = sum(f["total"] for f in fw_list)
                total_mapped = sum(f["mapped_controls"] for f in fw_list)
                ims_summary = {
                    "framework_count": len(fw_ids),
                    "total_mappings": total_mappings,
                    "total_controls": total_ctrls,
                    "mapped_controls": total_mapped,
                }
            except Exception:
                pass
    finally:
        db.close()
    return _aria_render(request, "frameworks.html", {
        "user": user, "module": "aria", "fw_list": fw_list,
        "ims_summary": ims_summary,
    }, active_section="frameworks")


# -- Framework Detail ---------------------------------------------------------

@router.get("/framework/{fw_id}", response_class=HTMLResponse)
@require_module("aria")
async def framework_detail(request: Request, fw_id: int,
                           status: str = "", priority: str = "",
                           category: str = "", search: str = ""):
    user = request.state.user
    db = get_db()
    try:
        fw = db.execute(
            "SELECT * FROM frameworks WHERE id=%s", (fw_id,)
        ).fetchone()
        if not fw:
            raise HTTPException(404)

        # Lazy-seed controls if none exist for this framework
        ctrl_count = db.execute(
            "SELECT COUNT(*) FROM controls WHERE framework_id=%s", (fw_id,)
        ).fetchone()[0]
        if ctrl_count == 0:
            try:
                from seeds.framework_controls import FRAMEWORK_CONTROLS
                from core.framework_service import bulk_create_controls
                seed_data = FRAMEWORK_CONTROLS.get(fw["name"])
                if seed_data:
                    db.close()
                    bulk_create_controls(fw_id, seed_data)
                    db = get_db()
            except Exception as exc:
                log.warning("Lazy-seed controls failed for fw %s: %s", fw_id, exc)

        q = ("SELECT c.*, "
             "(SELECT COUNT(*) FROM evidence_links el "
             " WHERE el.module='aria' AND el.entity_type='control' AND el.entity_id=c.id"
             ") AS evidence_count "
             "FROM controls c WHERE c.framework_id=%s")
        params = [fw_id]
        if status:
            q += " AND c.status=%s"
            params.append(status)
        if priority:
            q += " AND c.priority=%s"
            params.append(priority)
        if category:
            q += " AND c.category=%s"
            params.append(category)
        if search:
            q += " AND (c.name LIKE %s OR c.ref LIKE %s OR c.description LIKE %s)"
            params += ["%" + search + "%", "%" + search + "%", "%" + search + "%"]
        q += " ORDER BY c.ref"
        controls = db.execute(q, params).fetchall()

        cats = db.execute(
            "SELECT DISTINCT category FROM controls "
            "WHERE framework_id=%s ORDER BY category",
            (fw_id,),
        ).fetchall()
        stats = db.execute("""
            SELECT
                COUNT(*) as total,
                SUM(CASE WHEN status='Implemented' THEN 1 ELSE 0 END) as implemented,
                SUM(CASE WHEN status='Not Started' THEN 1 ELSE 0 END) as not_started,
                SUM(CASE WHEN status='In Progress' THEN 1 ELSE 0 END) as in_progress,
                SUM(CASE WHEN status='Approved' THEN 1 ELSE 0 END) as approved
            FROM controls WHERE framework_id=%s
        """, (fw_id,)).fetchone()
    finally:
        db.close()

    return _aria_render(request, "framework.html", {
        "user": user, "module": "aria", "fw": fw,
        "controls": controls, "stats": stats, "categories": cats,
        "filters": {
            "status": status, "priority": priority,
            "category": category, "search": search,
        },
    }, active_section="frameworks")


# -- Control Update -----------------------------------------------------------

@router.post("/control/{ctrl_id}/update")
@require_module("aria")
async def update_control(request: Request, ctrl_id: int,
                         status: str = Form(None), priority: str = Form(None),
                         owner: str = Form(None), target_date: str = Form(None),
                         review_date: str = Form(None), notes: str = Form(None),
                         document_title: str = Form(None),
                         version: str = Form(None),
                         evidence_ref: str = Form(None)):
    user = request.state.user
    db = get_db()
    try:
        old = db.execute(
            "SELECT * FROM controls WHERE id=%s", (ctrl_id,)
        ).fetchone()
        if not old:
            return JSONResponse({"error": "Control not found"}, 404)
        if not _can_edit_control(user, dict(old)):
            return JSONResponse(
                {"error": "You do not have permission to update this control."},
                403,
            )
        updates, params = [], []
        for field, val in [
            ("status", status), ("priority", priority), ("owner", owner),
            ("target_date", target_date), ("review_date", review_date),
            ("notes", notes), ("document_title", document_title),
            ("version", version), ("evidence_ref", evidence_ref),
        ]:
            if val is not None:
                updates.append(field + "=%s")
                params.append(val)
        if updates:
            updates.append("last_updated=%s")
            params.append(datetime.now().strftime("%Y-%m-%d"))
            params.append(ctrl_id)
            db.execute(
                "UPDATE controls SET " + ", ".join(updates) + " WHERE id=%s",
                params,
            )
            db.commit()
            old_status = old["status"] or ""
            log_audit(user, "aria", "Updated control to '" + str(status) + "'",
                      "control", ctrl_id, old_status)

            # Emit control updated event when status changes
            if status and status != old_status:
                emit(
                    ARIA_CONTROL_UPDATED,
                    source_module="aria",
                    entity_type="control",
                    entity_id=ctrl_id,
                    payload={
                        "ref": old["ref"],
                        "name": old["name"],
                        "old_status": old_status,
                        "new_status": status,
                        "framework_id": old["framework_id"],
                    },
                    user_id=user["id"],
                )
    finally:
        db.close()

    return JSONResponse({"ok": True})


# -- Stats API ----------------------------------------------------------------

@router.get("/api/frameworks")
@require_module("aria")
async def api_frameworks_list(request: Request):
    """Return active frameworks as JSON — used by cross-module integrations."""
    db = get_db()
    try:
        rows = db.execute(
            "SELECT id, name, color, description FROM frameworks WHERE is_active=1 ORDER BY name"
        ).fetchall()
    finally:
        db.close()
    return JSONResponse({"items": [dict(r) for r in rows]})


@router.get("/api/stats")
@require_module("aria")
async def api_stats(request: Request):
    user = request.state.user
    db = get_db()
    try:
        frameworks = db.execute("SELECT * FROM frameworks WHERE is_active = 1").fetchall()
        data = []
        for fw in frameworks:
            stats = db.execute("""
                SELECT
                    SUM(CASE WHEN status='Implemented' THEN 1 ELSE 0 END) as implemented,
                    SUM(CASE WHEN status='In Progress' THEN 1 ELSE 0 END) as in_progress,
                    SUM(CASE WHEN status='Not Started' THEN 1 ELSE 0 END) as not_started,
                    SUM(CASE WHEN status='Approved' THEN 1 ELSE 0 END) as approved,
                    SUM(CASE WHEN status='Draft' THEN 1 ELSE 0 END) as draft,
                    SUM(CASE WHEN status='Needs Update' THEN 1 ELSE 0 END) as needs_update,
                    COUNT(*) as total
                FROM controls WHERE framework_id=%s
            """, (fw["id"],)).fetchone()
            data.append({
                "name": fw["name"], "color": fw["color"],
                "implemented": stats["implemented"] or 0,
                "in_progress": stats["in_progress"] or 0,
                "not_started": stats["not_started"] or 0,
                "approved": stats["approved"] or 0,
                "draft": stats["draft"] or 0,
                "needs_update": stats["needs_update"] or 0,
                "total": stats["total"] or 0,
            })
    finally:
        db.close()
    return JSONResponse(data)


@router.get("/api/dashboard")
@require_module("aria")
async def api_dashboard(request: Request):
    """Dashboard summary used by the live-refresh poll every 5 minutes."""
    db = get_db()
    try:
        row = db.execute(
            "SELECT COUNT(*) as total, "
            "SUM(CASE WHEN status IN ('Implemented','Approved') THEN 1 ELSE 0 END) as compliant "
            "FROM controls c JOIN frameworks f ON f.id=c.framework_id WHERE f.is_active=1"
        ).fetchone()
        total = row["total"] or 0
        compliant = row["compliant"] or 0
        overall_pct = round(compliant / total * 100) if total else 0
    finally:
        db.close()
    return JSONResponse({"totals": {"overall_pct": overall_pct, "total": total, "compliant": compliant}})


@router.get("/api/controls/{fw_id}")
@require_module("aria")
async def api_controls(request: Request, fw_id: int):
    user = request.state.user
    db = get_db()
    try:
        controls = db.execute(
            "SELECT id, ref, name, description, doc_type, status "
            "FROM controls WHERE framework_id=%s ORDER BY ref",
            (fw_id,),
        ).fetchall()
    finally:
        db.close()

    # Lazy-seed: if no controls found, try to populate from seed data
    if not controls:
        try:
            from core.framework_service import get_framework, bulk_create_controls
            fw = get_framework(fw_id)
            if fw:
                from seeds.framework_controls import FRAMEWORK_CONTROLS
                seed_data = FRAMEWORK_CONTROLS.get(fw["name"])
                if seed_data:
                    bulk_create_controls(fw_id, seed_data)
                    db2 = get_db()
                    try:
                        controls = db2.execute(
                            "SELECT id, ref, name, description, doc_type, status "
                            "FROM controls WHERE framework_id=%s ORDER BY ref",
                            (fw_id,),
                        ).fetchall()
                    finally:
                        db2.close()
        except Exception as exc:
            log.warning("Lazy-seed controls failed for fw %s: %s", fw_id, exc)

    return JSONResponse([dict(c) for c in controls])


@router.post("/api/force-seed-controls/{fw_id}")
@require_module("aria")
async def api_force_seed_controls(request: Request, fw_id: int):
    """Force-seed controls for a specific framework."""
    db = get_db()
    try:
        fw = db.execute(
            "SELECT id, name FROM frameworks WHERE id=%s", (fw_id,)
        ).fetchone()
        if not fw:
            return JSONResponse({"error": "Framework not found", "seeded": 0}, 404)

        # Check existing count
        existing = db.execute(
            "SELECT COUNT(*) as c FROM controls WHERE framework_id=%s",
            (fw_id,)
        ).fetchone()["c"]

        if existing > 0:
            return JSONResponse({
                "message": f"Framework already has {existing} controls",
                "seeded": 0,
                "existing": existing,
            })

        from seeds.framework_controls import FRAMEWORK_CONTROLS
        seed_data = FRAMEWORK_CONTROLS.get(fw["name"])
        if not seed_data:
            # Try partial match
            fw_name_lower = fw["name"].lower()
            for key, data in FRAMEWORK_CONTROLS.items():
                if key.lower() in fw_name_lower or fw_name_lower in key.lower():
                    seed_data = data
                    log.info("Force-seed: matched '%s' to seed key '%s'",
                             fw["name"], key)
                    break

        if not seed_data:
            return JSONResponse({
                "error": f"No seed data for '{fw['name']}'",
                "seeded": 0,
                "available_keys": list(FRAMEWORK_CONTROLS.keys()),
            })

        count = 0
        for ctrl in seed_data:
            try:
                db.execute(
                    "INSERT INTO controls "
                    "(framework_id, ref, name, description, category, "
                    "doc_type, priority, last_updated) "
                    "VALUES (%s, %s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING",
                    (
                        fw_id,
                        ctrl.get("ref", ""),
                        ctrl.get("name", ""),
                        ctrl.get("description", ""),
                        ctrl.get("category", ""),
                        ctrl.get("doc_type", "Policy"),
                        ctrl.get("priority", "High"),
                        utcnow().isoformat(),
                    ),
                )
                count += 1
            except Exception as e:
                log.warning("Force-seed insert error: %s", e)
        db.commit()

        # Update total
        final_count = db.execute(
            "SELECT COUNT(*) as c FROM controls WHERE framework_id=%s",
            (fw_id,)
        ).fetchone()["c"]
        db.execute(
            "UPDATE frameworks SET total_controls=%s WHERE id=%s",
            (final_count, fw_id),
        )
        db.commit()

        log.info("Force-seeded %d controls for '%s' (fw_id=%d)",
                 final_count, fw["name"], fw_id)
        return JSONResponse({
            "message": f"Seeded {final_count} controls for {fw['name']}",
            "seeded": final_count,
        })
    except Exception as exc:
        log.exception("Force-seed failed: %s", exc)
        return JSONResponse({"error": str(exc), "seeded": 0}, 500)
    finally:
        db.close()


@router.get("/api/debug/controls-status")
@require_module("aria")
async def api_debug_controls_status(request: Request):
    """Diagnostic endpoint: shows frameworks and control counts."""
    db = get_db()
    try:
        frameworks = db.execute(
            "SELECT id, name, is_active FROM frameworks ORDER BY name"
        ).fetchall()
        result = []
        for fw in frameworks:
            count = db.execute(
                "SELECT COUNT(*) FROM controls WHERE framework_id=%s",
                (fw["id"],),
            ).fetchone()[0]
            result.append({
                "id": fw["id"], "name": fw["name"],
                "is_active": fw["is_active"], "control_count": count,
            })
        total = db.execute("SELECT COUNT(*) FROM controls").fetchone()[0]
        return JSONResponse({
            "total_controls": total,
            "frameworks": result,
        })
    finally:
        db.close()


# -- Documents ----------------------------------------------------------------

@router.get("/documents", response_class=HTMLResponse)
@require_module("aria")
async def documents_page(request: Request,
                         framework: str = "", status: str = "",
                         doc_type: str = "", search: str = ""):
    user = request.state.user
    db = get_db()
    try:
        q = ("SELECT id, doc_id, framework, control_ref, title, doc_type, "
             "version, status, owner, approver, effective_date, review_date, "
             "location, comments, created_at, updated_at "
             "FROM aria_documents WHERE 1=1")
        params = []
        if framework:
            q += " AND framework LIKE %s"
            params.append("%" + framework + "%")
        if status:
            q += " AND status=%s"
            params.append(status)
        if doc_type:
            q += " AND doc_type=%s"
            params.append(doc_type)
        if search:
            q += " AND (title LIKE %s OR control_ref LIKE %s OR owner LIKE %s)"
            params += ["%" + search + "%"] * 3
        q += " ORDER BY framework, doc_id"
        docs = [dict(r) for r in db.execute(q, params).fetchall()]

        frameworks = db.execute(
            "SELECT id, name FROM frameworks WHERE is_active = 1 ORDER BY name"
        ).fetchall()

        total = db.execute("SELECT COUNT(*) FROM aria_documents").fetchone()[0]
        approved = db.execute(
            "SELECT COUNT(*) FROM aria_documents WHERE status='Approved'"
        ).fetchone()[0]
        draft = db.execute(
            "SELECT COUNT(*) FROM aria_documents WHERE status='Draft'"
        ).fetchone()[0]
        ai_gen = db.execute(
            "SELECT COUNT(*) FROM aria_documents "
            "WHERE comments LIKE '%AI Generated%'"
        ).fetchone()[0]
        in_30 = (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d")
        review_due = db.execute(
            "SELECT COUNT(*) FROM aria_documents "
            "WHERE review_date IS NOT NULL AND review_date != '' "
            "AND status != 'Retired' AND review_date <= %s",
            (in_30,),
        ).fetchone()[0]
    finally:
        db.close()

    return _aria_render(request, "documents.html", {
        "user": user, "module": "aria", "docs": docs,
        "frameworks": frameworks,
        "filters": {
            "framework": framework, "status": status,
            "doc_type": doc_type, "search": search,
        },
        "stats": {
            "total": total, "approved": approved,
            "draft": draft, "ai_gen": ai_gen, "review_due": review_due,
        },
    }, active_section="documents")


@router.post("/documents/add")
@require_module("aria")
async def add_document(request: Request,
                       framework: str = Form(...),
                       control_ref: str = Form(""),
                       title: str = Form(...),
                       doc_type: str = Form("Policy"),
                       version: str = Form("1.0"),
                       status: str = Form("Draft"),
                       owner: str = Form(""),
                       approver: str = Form(""),
                       effective_date: str = Form(""),
                       review_date: str = Form(""),
                       location: str = Form(""),
                       comments: str = Form("")):
    user = request.state.user
    if not has_capability(user, "aria.policy.create"):
        return JSONResponse(
            {"error": "You need policy author or compliance manager role."},
            403,
        )
    if status == "Approved" and not has_capability(user, "aria.policy.approve"):
        status = "Draft"

    db = get_db()
    try:
        count = db.execute("SELECT COUNT(*) FROM aria_documents").fetchone()[0]
        doc_id = "DOC-%04d" % (count + 1)
        now = datetime.now().isoformat()
        new_id = insert_returning_id(db,"""
            INSERT INTO aria_documents
            (doc_id, framework, control_ref, title, doc_type, version, status,
             owner, approver, effective_date, review_date, location, comments,
             created_at, updated_at)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        """, (doc_id, framework, control_ref, title, doc_type, version, status,
              owner, approver, effective_date or None, review_date or None,
              location, comments, now, now))
        db.commit()

        if review_date:
            _schedule_review_reminder(db, new_id, doc_id, title, review_date, user)
            db.commit()

        log_audit(user, "aria", "Added document " + doc_id + ": " + title,
                  "document", new_id, doc_id)

        # Emit policy published event when created as Approved
        if status == "Approved":
            emit(
                ARIA_POLICY_PUBLISHED,
                source_module="aria",
                entity_type="document",
                entity_id=new_id,
                payload={
                    "doc_id": doc_id, "title": title,
                    "doc_type": doc_type, "framework": framework,
                    "control_ref": control_ref, "version": version,
                },
                user_id=user["id"],
            )
    finally:
        db.close()

    return JSONResponse({"ok": True, "doc_id": doc_id})


@router.post("/documents/parse-upload")
@require_module("aria")
async def parse_document_upload(request: Request, file: UploadFile = File(...)):
    """Extract metadata from an uploaded file without saving anything."""
    content = await file.read()
    if len(content) > ARIA_MAX_FILE:
        return JSONResponse({"error": "File too large (max 50 MB)"}, 413)
    ext = Path(file.filename or "").suffix.lower()
    if ext not in _ALLOWED_DOC_EXT:
        return JSONResponse({"error": f"File type '{ext}' not allowed"}, 400)
    meta = _extract_document_metadata(content, ext, file.filename or "untitled")
    return JSONResponse({"ok": True, **meta})


@router.post("/documents/upload-new")
@require_module("aria")
async def upload_new_document(
    request: Request,
    file: UploadFile = File(None),
    framework: str = Form(...),
    control_ref: str = Form(""),
    title: str = Form(...),
    doc_type: str = Form("Policy"),
    version: str = Form("1.0"),
    status: str = Form("Draft"),
    owner: str = Form(""),
    approver: str = Form(""),
    effective_date: str = Form(""),
    review_date: str = Form(""),
    location: str = Form(""),
    comments: str = Form(""),
):
    """Create a new document record from an uploaded file plus metadata."""
    user = request.state.user
    if not has_capability(user, "aria.policy.create"):
        return JSONResponse(
            {"error": "You need policy author or compliance manager role."}, 403
        )
    if status == "Approved" and not has_capability(user, "aria.policy.approve"):
        status = "Draft"

    db = get_db()
    try:
        count = db.execute("SELECT COUNT(*) FROM aria_documents").fetchone()[0]
        doc_id = "DOC-%04d" % (count + 1)
        now = datetime.now().isoformat()

        stored_name = None
        orig_filename = None
        file_size = 0
        if file and file.filename:
            content = await file.read()
            if len(content) > ARIA_MAX_FILE:
                return JSONResponse({"error": "File too large (max 50 MB)"}, 413)
            ext = Path(file.filename).suffix.lower()
            if ext not in _ALLOWED_DOC_EXT:
                return JSONResponse({"error": f"File type '{ext}' not allowed"}, 400)
            ARIA_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
            stored_name = f"{uuid.uuid4().hex}{ext}"
            (ARIA_UPLOAD_DIR / stored_name).write_bytes(content)
            orig_filename = file.filename
            file_size = len(content)

        source_tag = "Uploaded" if stored_name else "Manual"
        comments_tagged = (
            (comments + "\n[" + source_tag + "]").strip() if comments
            else "[" + source_tag + "]"
        )

        new_id = insert_returning_id(db, """
            INSERT INTO aria_documents
            (doc_id, framework, control_ref, title, doc_type, version, status,
             owner, approver, effective_date, review_date, location, comments,
             file_path, file_name, file_size, created_at, updated_at)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        """, (
            doc_id, framework, control_ref, title, doc_type, version, status,
            owner, approver, effective_date or None, review_date or None,
            location, comments_tagged,
            stored_name, orig_filename, file_size, now, now,
        ))
        db.commit()

        if review_date:
            _schedule_review_reminder(db, new_id, doc_id, title, review_date, user)
            db.commit()

        log_audit(user, "aria", f"Added document {doc_id}: {title}",
                  "document", new_id, doc_id)

        if status == "Approved":
            emit(
                ARIA_POLICY_PUBLISHED,
                source_module="aria", entity_type="document", entity_id=new_id,
                payload={
                    "doc_id": doc_id, "title": title, "doc_type": doc_type,
                    "framework": framework, "version": version,
                },
                user_id=user["id"],
            )
    finally:
        db.close()

    return JSONResponse({"ok": True, "doc_id": doc_id})


@router.post("/documents/update/{doc_id}")
@require_module("aria")
async def update_document(request: Request, doc_id: str,
                          title: str = Form(None),
                          status: str = Form(None),
                          version: str = Form(None),
                          owner: str = Form(None),
                          approver: str = Form(None),
                          effective_date: str = Form(None),
                          review_date: str = Form(None),
                          location: str = Form(None),
                          comments: str = Form(None)):
    user = request.state.user
    db = get_db()
    try:
        existing = db.execute(
            "SELECT * FROM aria_documents WHERE doc_id=%s", (doc_id,)
        ).fetchone()
        if not existing:
            return JSONResponse({"error": "Document not found"}, 404)
        doc = dict(existing)

        is_own = (doc.get("owner") or "").strip().lower() in (
            (user.get("full_name") or "").strip().lower(),
            (user.get("username") or "").strip().lower(),
        )
        can_edit = has_capability(user, "aria.policy.edit_any") or (
            has_capability(user, "aria.policy.edit_own") and is_own
        )
        if not can_edit:
            return JSONResponse(
                {"error": "You do not have permission to edit this document."},
                403,
            )

        if status == "Approved" and status != doc.get("status"):
            if not _can_approve_policy(user, doc):
                return JSONResponse({
                    "error": "You cannot approve this document. "
                             "Policy authors cannot approve their own drafts."
                }, 403)

        updates, params = [], []
        for field, val in [
            ("title", title), ("status", status), ("version", version),
            ("owner", owner), ("approver", approver),
            ("effective_date", effective_date), ("review_date", review_date),
            ("location", location), ("comments", comments),
        ]:
            if val is not None:
                updates.append(field + "=%s")
                params.append(val)
        if updates:
            updates.append("updated_at=%s")
            params.append(datetime.now().isoformat())
            params.append(doc_id)
            db.execute(
                "UPDATE aria_documents SET " + ", ".join(updates) + " WHERE doc_id=%s",
                params,
            )
            db.commit()

            if review_date and review_date != (doc.get("review_date") or ""):
                _schedule_review_reminder(
                    db, doc["id"], doc_id,
                    title or doc.get("title", ""),
                    review_date, user,
                )
                db.commit()

            log_audit(user, "aria", "Updated document " + doc_id,
                      "document", doc["id"], doc_id)

            # Emit policy events on status change
            old_status = doc.get("status", "")
            if status and status != old_status:
                evt = ARIA_POLICY_PUBLISHED if status == "Approved" else ARIA_POLICY_UPDATED
                emit(
                    evt,
                    source_module="aria",
                    entity_type="document",
                    entity_id=doc.get("id", 0),
                    payload={
                        "doc_id": doc_id,
                        "title": title or doc.get("title", ""),
                        "old_status": old_status,
                        "new_status": status,
                        "version": version or doc.get("version", ""),
                    },
                    user_id=user["id"],
                )
    finally:
        db.close()

    return JSONResponse({"ok": True})


@router.post("/documents/delete/{doc_id}")
@require_module("aria")
async def delete_document(request: Request, doc_id: str):
    user = request.state.user
    if not has_capability(user, "aria.policy.delete"):
        return JSONResponse(
            {"error": "Only System Administrators can delete documents."},
            403,
        )
    db = get_db()
    try:
        db.execute("DELETE FROM aria_documents WHERE doc_id=%s", (doc_id,))
        db.commit()
        log_audit(user, "aria", "Deleted document " + doc_id,
                  "document", 0, doc_id)
    finally:
        db.close()
    return JSONResponse({"ok": True})


# -- Document Revision Upload + Template Application -------------------------

import os
import uuid
import hashlib
from pathlib import Path

ARIA_UPLOAD_DIR = Path(os.getenv("ARIA_UPLOAD_DIR", "data/aria_uploads"))
ARIA_TEMPLATE_DIR = Path(os.getenv("ARIA_TEMPLATE_DIR", "data/aria_templates"))
ARIA_MAX_FILE = 50 * 1024 * 1024  # 50 MB

_ALLOWED_DOC_EXT = frozenset({".docx", ".doc", ".pdf", ".odt", ".xlsx", ".pptx"})

# ── Document metadata extraction helpers ──────────────────────────────────────

_MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


def _parse_date_str(raw: str) -> str:
    """Best-effort conversion of a human date string to YYYY-MM-DD."""
    raw = raw.strip()
    m = _re.match(r"(\d{4})-(\d{1,2})-(\d{1,2})", raw)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    m = _re.match(r"(\d{1,2})[/.-](\d{1,2})[/.-](\d{4})", raw)
    if m:
        a, b, y = int(m.group(1)), int(m.group(2)), m.group(3)
        if a > 12:
            return f"{y}-{b:02d}-{a:02d}"
        return f"{y}-{a:02d}-{b:02d}"
    m = _re.match(r"(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})", raw)
    if m:
        day, mon, yr = int(m.group(1)), m.group(2)[:3].lower(), m.group(3)
        if mon in _MONTH_MAP:
            return f"{yr}-{_MONTH_MAP[mon]:02d}-{day:02d}"
    m = _re.match(r"([A-Za-z]+)\s+(\d{4})", raw)
    if m:
        mon, yr = m.group(1)[:3].lower(), m.group(2)
        if mon in _MONTH_MAP:
            return f"{yr}-{_MONTH_MAP[mon]:02d}-01"
    return ""


def _scan_text_for_metadata(text: str) -> dict:
    """Scan document body text for common policy metadata patterns."""
    result: dict = {}
    _date_pat = (
        r"(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}"
        r"|\d{4}[/-]\d{2}[/-]\d{2}"
        r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}"
        r"|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4})"
    )
    m = _re.search(r"[Vv]ersion[:\s]+(\d+\.\d+(?:\.\d+)?)", text)
    if m:
        result["version"] = m.group(1)

    m = _re.search(r"(?:Next\s+)?[Rr]eview\s+[Dd]ate[:\s]+" + _date_pat, text)
    if m:
        parsed = _parse_date_str(m.group(1))
        if parsed:
            result["review_date"] = parsed

    m = _re.search(r"[Ee]ffective\s+(?:[Dd]ate|[Ff]rom)[:\s]+" + _date_pat, text)
    if m:
        parsed = _parse_date_str(m.group(1))
        if parsed:
            result["effective_date"] = parsed

    m = _re.search(
        r"(?:[Dd]ocument\s+)?[Oo]wner[:\s]+([A-Z][a-zA-Z ]{2,40}?)(?:\n|,|\|)", text
    )
    if not m:
        m = _re.search(r"[Pp]repared\s+[Bb]y[:\s]+([A-Z][a-zA-Z ]{2,40}?)(?:\n|,|\|)", text)
    if m:
        result["owner"] = m.group(1).strip()

    m = _re.search(r"[Aa]pproved?\s+[Bb]y[:\s]+([A-Z][a-zA-Z ]{2,40}?)(?:\n|,|\|)", text)
    if m:
        result["approver"] = m.group(1).strip()

    for dtype in ("Procedure", "Standard", "Guideline", "Template", "Record", "Plan"):
        if _re.search(r"\b" + dtype + r"\b", text[:800]):
            result["doc_type"] = dtype
            break

    return result


def _extract_document_metadata(content: bytes, ext: str, filename: str) -> dict:
    """Extract metadata from a document file without saving it."""
    stem = Path(filename).stem.replace("_", " ").replace("-", " ")
    result: dict = {
        "title": stem, "version": "", "review_date": "",
        "effective_date": "", "owner": "", "approver": "", "doc_type": "Policy",
    }
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            props = doc.core_properties
            if props.title:
                result["title"] = props.title
            if props.version:
                result["version"] = props.version
            if props.last_modified_by:
                result["owner"] = result["owner"] or props.last_modified_by
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            result.update(_scan_text_for_metadata(full_text))
        except Exception as exc:
            log.debug("DOCX metadata extract failed: %s", exc)
    elif ext == ".pdf":
        try:
            raw = content.decode("latin-1", errors="ignore")
            for pdf_key, field in [("/Title", "title"), ("/Author", "owner")]:
                m = _re.search(_re.escape(pdf_key) + r"\s*\(([^)]{1,200})\)", raw)
                if m and m.group(1).strip():
                    result[field] = result[field] or m.group(1).strip()
            text_chunks = _re.findall(r"\(([^)]{4,200})\)", raw)
            combined = " ".join(text_chunks[:300])
            result.update({k: v for k, v in _scan_text_for_metadata(combined).items() if v})
        except Exception as exc:
            log.debug("PDF metadata extract failed: %s", exc)
    return result


def _schedule_review_reminder(db, doc_int_id: int, doc_id: str, title: str,
                               review_date_str: str, user: dict) -> None:
    """Insert 30-day and 7-day ahead email reminders for a document review date."""
    from datetime import date
    if not review_date_str:
        return
    try:
        rd = date.fromisoformat(review_date_str)
    except (ValueError, TypeError):
        return
    today = date.today()
    if rd <= today:
        return
    email = (user.get("email") or "").strip()
    if not email:
        try:
            row = db.execute(
                "SELECT email FROM users WHERE id=%s", (user["id"],)
            ).fetchone()
            if row:
                email = (row["email"] or "").strip()
        except Exception:
            pass
    if not email:
        return
    from datetime import date, timedelta
    for days_before in (30, 7):
        remind_date = rd - timedelta(days=days_before)
        if remind_date <= today:
            continue
        remind_at = remind_date.isoformat() + " 09:00:00"
        existing = db.execute(
            "SELECT id FROM email_reminders "
            "WHERE module='aria' AND entity_type='document' AND entity_id=%s "
            "AND is_sent=0 AND remind_at=%s",
            (doc_int_id, remind_at),
        ).fetchone()
        if existing:
            continue
        db.execute(
            "INSERT INTO email_reminders "
            "(module, entity_type, entity_id, title, message, "
            " recipient_email, remind_at, created_by) "
            "VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (
                "aria", "document", doc_int_id,
                f"Policy Review Due: {title}",
                (
                    f"The document '{title}' ({doc_id}) is due for review in "
                    f"{days_before} days, on {review_date_str}. "
                    f"Please log in to ThemisIQ to review and update the document."
                ),
                email, remind_at, user["id"],
            ),
        )


@router.post("/documents/{doc_id}/upload-revision")
@require_module("aria")
async def upload_document_revision(request: Request, doc_id: str,
                                    file: UploadFile = File(...),
                                    notes: str = Form("")):
    """Upload a revised/edited document file to replace the AI-generated draft."""
    user = request.state.user
    if not has_capability(user, "aria.policy.edit_own") and not has_capability(user, "aria.policy.edit_any"):
        return JSONResponse({"error": "Permission denied"}, 403)

    content = await file.read()
    if len(content) > ARIA_MAX_FILE:
        return JSONResponse({"error": "File too large (max 50 MB)"}, 413)

    ext = Path(file.filename or "").suffix.lower()
    if ext not in _ALLOWED_DOC_EXT:
        return JSONResponse({"error": f"File type '{ext}' not allowed"}, 400)

    db = get_db()
    try:
        doc = db.execute("SELECT * FROM aria_documents WHERE doc_id=%s", (doc_id,)).fetchone()
        if not doc:
            return JSONResponse({"error": "Document not found"}, 404)
        doc = dict(doc)

        ARIA_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        stored_name = f"{uuid.uuid4().hex}{ext}"
        dest = ARIA_UPLOAD_DIR / stored_name
        dest.write_bytes(content)

        # Save current body as AI draft if this is the first upload
        if doc.get("body") and not doc.get("ai_draft_body"):
            db.execute(
                "UPDATE aria_documents SET ai_draft_body=%s WHERE doc_id=%s",
                (doc["body"], doc_id),
            )

        # Create revision record
        old_version = doc.get("version", "1.0")
        db.execute(
            "INSERT INTO aria_doc_revisions "
            "(document_id, version, file_path, file_name, file_size, "
            " revision_type, notes, uploaded_by) "
            "VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (doc["id"], old_version, doc.get("file_path", ""),
             doc.get("file_name", ""), doc.get("file_size", 0),
             "superseded", f"Replaced by upload", user["id"]),
        )

        # Bump version
        try:
            parts = old_version.split(".")
            parts[-1] = str(int(parts[-1]) + 1)
            new_version = ".".join(parts)
        except (ValueError, IndexError):
            new_version = old_version + ".1"

        # Update document record
        db.execute(
            "UPDATE aria_documents SET "
            "file_path=%s, file_name=%s, file_size=%s, version=%s, "
            "status='Under Review', updated_at=CURRENT_TIMESTAMP "
            "WHERE doc_id=%s",
            (str(stored_name), file.filename, len(content),
             new_version, doc_id),
        )
        db.commit()
        log_audit(user, "aria", f"Uploaded revision for {doc_id}: {file.filename}",
                  "document", doc["id"], doc_id)

        # Emit status change event
        if doc.get("status") != "Under Review":
            emit(
                ARIA_POLICY_UPDATED,
                source_module="aria", entity_type="document",
                entity_id=doc["id"],
                payload={
                    "doc_id": doc_id, "title": doc.get("title", ""),
                    "old_status": doc.get("status", ""),
                    "new_status": "Under Review",
                    "version": new_version,
                },
                user_id=user["id"],
            )
    finally:
        db.close()

    return JSONResponse({
        "success": True, "version": new_version,
        "file_name": file.filename,
    })


@router.get("/documents/{doc_id}/download")
@require_module("aria")
async def download_document(request: Request, doc_id: str):
    """Download the current document file (uploaded revision or branded version)."""
    from fastapi.responses import FileResponse as FR
    db = get_db()
    try:
        doc = db.execute(
            "SELECT file_path, file_name, branded_file_path, title "
            "FROM aria_documents WHERE doc_id=%s", (doc_id,)
        ).fetchone()
        if not doc:
            raise HTTPException(404, "Document not found")

        # Prefer branded version, fall back to uploaded file
        fp_str = doc["branded_file_path"] or doc["file_path"]
        if not fp_str:
            raise HTTPException(404, "No file uploaded for this document yet")

        fp = (ARIA_UPLOAD_DIR / fp_str).resolve()
        if not str(fp).startswith(str(ARIA_UPLOAD_DIR.resolve())):
            raise HTTPException(403, "Access denied")
        if not fp.exists():
            raise HTTPException(404, "File not found on disk")

        fname = doc["file_name"] or doc["title"] or "document"
    finally:
        db.close()

    return FR(str(fp), filename=fname, media_type="application/octet-stream")


@router.get("/documents/{doc_id}/revisions")
@require_module("aria")
async def document_revisions(request: Request, doc_id: str):
    """Get revision history for a document."""
    db = get_db()
    try:
        doc = db.execute("SELECT id FROM aria_documents WHERE doc_id=%s", (doc_id,)).fetchone()
        if not doc:
            return JSONResponse({"error": "Document not found"}, 404)
        rows = db.execute(
            "SELECT r.*, u.full_name AS uploaded_by_name "
            "FROM aria_doc_revisions r "
            "LEFT JOIN users u ON r.uploaded_by=u.id "
            "WHERE r.document_id=%s ORDER BY r.created_at DESC",
            (doc[0],),
        ).fetchall()
    finally:
        db.close()
    return JSONResponse([dict(r) for r in rows])


# -- Document Templates (Branding) -------------------------------------------

@router.get("/templates", response_class=HTMLResponse)
@require_module("aria")
async def templates_page(request: Request):
    """Template management page."""
    user = request.state.user
    db = get_db()
    try:
        tpls = [dict(r) for r in db.execute(
            "SELECT t.*, u.full_name AS created_by_name "
            "FROM aria_doc_templates t "
            "LEFT JOIN users u ON t.created_by=u.id "
            "ORDER BY t.is_default DESC, t.name"
        ).fetchall()]
    finally:
        db.close()
    return _aria_render(request, "templates.html", {
        "user": user, "module": "aria", "templates": tpls,
    }, active_section="templates")


@router.get("/api/templates")
@require_module("aria")
async def api_templates_list(request: Request):
    """List all document templates."""
    db = get_db()
    try:
        rows = db.execute(
            "SELECT t.*, u.full_name AS created_by_name "
            "FROM aria_doc_templates t "
            "LEFT JOIN users u ON t.created_by=u.id "
            "ORDER BY t.is_default DESC, t.name"
        ).fetchall()
    finally:
        db.close()
    return JSONResponse([dict(r) for r in rows])


@router.post("/api/templates")
@require_capability("aria.policy.create")
async def api_templates_upload(request: Request,
                                file: UploadFile = File(...),
                                name: str = Form(...),
                                description: str = Form(""),
                                doc_type: str = Form("Policy"),
                                is_default: str = Form("0")):
    """Upload a new branding template (.docx)."""
    user = request.state.user
    content = await file.read()
    if len(content) > ARIA_MAX_FILE:
        return JSONResponse({"error": "File too large"}, 413)

    ext = Path(file.filename or "").suffix.lower()
    if ext != ".docx":
        return JSONResponse({"error": "Templates must be .docx files"}, 400)

    ARIA_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    stored = f"{uuid.uuid4().hex}.docx"
    (ARIA_TEMPLATE_DIR / stored).write_bytes(content)

    db = get_db()
    try:
        # If setting as default, clear other defaults for this doc_type
        if is_default == "1":
            db.execute(
                "UPDATE aria_doc_templates SET is_default=0 WHERE doc_type=%s",
                (doc_type,),
            )
        tid = insert_returning_id(db,
            "INSERT INTO aria_doc_templates "
            "(name, description, doc_type, file_path, file_name, file_size, "
            " is_default, created_by) "
            "VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (name.strip(), description.strip(), doc_type, stored,
             file.filename, len(content), int(is_default == "1"),
             user["id"]),
        )
        db.commit()
        log_audit(user, "aria", f"Uploaded template: {name}", "template", tid)
    finally:
        db.close()
    return JSONResponse({"ok": True, "id": tid}, status_code=201)


@router.delete("/api/templates/{tid}")
@require_capability("aria.policy.delete")
async def api_templates_delete(request: Request, tid: int):
    """Delete a branding template."""
    user = request.state.user
    db = get_db()
    try:
        row = db.execute("SELECT file_path FROM aria_doc_templates WHERE id=%s", (tid,)).fetchone()
        if row and row[0]:
            fp = ARIA_TEMPLATE_DIR / row[0]
            try:
                if fp.exists():
                    fp.unlink()
            except OSError:
                pass
        db.execute("DELETE FROM aria_doc_templates WHERE id=%s", (tid,))
        db.commit()
        log_audit(user, "aria", f"Deleted template #{tid}", "template", tid)
    finally:
        db.close()
    return JSONResponse({"ok": True})


@router.get("/api/templates/{tid}/download")
@require_module("aria")
async def api_templates_download(request: Request, tid: int):
    """Download a template file."""
    from fastapi.responses import FileResponse as FR
    db = get_db()
    try:
        row = db.execute(
            "SELECT file_path, file_name FROM aria_doc_templates WHERE id=%s", (tid,)
        ).fetchone()
        if not row:
            raise HTTPException(404, "Template not found")
    finally:
        db.close()
    fp = (ARIA_TEMPLATE_DIR / row["file_path"]).resolve()
    if not str(fp).startswith(str(ARIA_TEMPLATE_DIR.resolve())):
        raise HTTPException(403, "Access denied")
    if not fp.exists():
        raise HTTPException(404, "Template file missing")
    return FR(str(fp), filename=row["file_name"],
              media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# -- Apply Branding Template to Document ------------------------------------

@router.post("/documents/{doc_id}/apply-template")
@require_capability("aria.policy.edit_own")
async def apply_template_to_document(request: Request, doc_id: str,
                                      template_id: int = Form(...)):
    """Apply a branding template to the uploaded document revision."""
    user = request.state.user
    db = get_db()
    try:
        doc = db.execute("SELECT * FROM aria_documents WHERE doc_id=%s", (doc_id,)).fetchone()
        if not doc:
            return JSONResponse({"error": "Document not found"}, 404)
        doc = dict(doc)

        if not doc.get("file_path"):
            return JSONResponse({"error": "No uploaded file to apply template to"}, 400)

        tpl = db.execute("SELECT * FROM aria_doc_templates WHERE id=%s", (template_id,)).fetchone()
        if not tpl:
            return JSONResponse({"error": "Template not found"}, 404)
        tpl = dict(tpl)

        # Apply template using branding engine
        from modules.aria.branding_engine import apply_template
        source_path = ARIA_UPLOAD_DIR / doc["file_path"]
        template_path = ARIA_TEMPLATE_DIR / tpl["file_path"]

        if not source_path.exists():
            return JSONResponse({"error": "Source document file missing"}, 404)
        if not template_path.exists():
            return JSONResponse({"error": "Template file missing"}, 404)

        branded_name = f"branded_{uuid.uuid4().hex}.docx"
        branded_path = ARIA_UPLOAD_DIR / branded_name

        apply_template(
            source_path=str(source_path),
            template_path=str(template_path),
            output_path=str(branded_path),
            logo_path=tpl.get("logo_path"),
            doc_title=doc.get("title", ""),
            doc_id=doc_id,
            version=doc.get("version", "1.0"),
            framework=doc.get("framework", ""),
        )

        # Update document record
        db.execute(
            "UPDATE aria_documents SET "
            "branded_file_path=%s, template_id=%s, updated_at=CURRENT_TIMESTAMP "
            "WHERE doc_id=%s",
            (branded_name, template_id, doc_id),
        )
        db.commit()
        log_audit(user, "aria",
                  f"Applied template '{tpl['name']}' to {doc_id}",
                  "document", doc["id"], doc_id)
    finally:
        db.close()

    return JSONResponse({"ok": True, "branded_file": branded_name})


# -- Risks --------------------------------------------------------------------

@router.get("/risks", response_class=HTMLResponse)
@require_module("aria")
async def risks_page(request: Request):
    user = request.state.user
    db = get_db()
    try:
        risks = db.execute(
            "SELECT * FROM aria_risks ORDER BY (likelihood * impact) DESC"
        ).fetchall()
    finally:
        db.close()
    return _aria_render(request, "risks.html", {
        "user": user, "module": "aria", "risks": risks,
    }, active_section="risks")


@router.post("/risks/add")
@require_module("aria")
async def add_risk(request: Request,
                   framework: str = Form(...),
                   control_ref: str = Form(""),
                   description: str = Form(...),
                   category: str = Form(""),
                   likelihood: int = Form(3),
                   impact: int = Form(3),
                   owner: str = Form(""),
                   mitigation: str = Form("")):
    user = request.state.user
    if not has_capability(user, "aria.risk.add"):
        return JSONResponse(
            {"error": "You do not have permission to add risks."}, 403
        )
    db = get_db()
    try:
        count = db.execute("SELECT COUNT(*) FROM aria_risks").fetchone()[0]
        risk_id = "RISK-%04d" % (count + 1)
        new_id = insert_returning_id(db,"""
            INSERT INTO aria_risks
            (risk_id, framework, control_ref, description, category,
             likelihood, impact, owner, mitigation, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 'Open')
        """, (risk_id, framework, control_ref, description, category,
              likelihood, impact, owner, mitigation))
        db.commit()
        log_audit(user, "aria", "Added risk " + risk_id,
                  "risk", 0, risk_id)

        # Emit risk created event
        emit(
            ARIA_RISK_CREATED,
            source_module="aria",
            entity_type="risk",
            entity_id=new_id,
            payload={
                "risk_id": risk_id, "description": description,
                "category": category, "framework": framework,
                "control_ref": control_ref,
                "likelihood": likelihood, "impact": impact,
                "risk_score": likelihood * impact,
            },
            user_id=user["id"],
        )
    finally:
        db.close()
    return JSONResponse({"ok": True, "risk_id": risk_id})


# -- Risk Update (status/owner/mitigation) -----------------------------------

@router.post("/risks/update/{risk_id}")
@require_module("aria")
async def update_risk(request: Request, risk_id: int,
                      status: str = Form(None),
                      owner: str = Form(None),
                      mitigation: str = Form(None),
                      likelihood: int = Form(None),
                      impact: int = Form(None)):
    user = request.state.user
    if not has_capability(user, "aria.risk.add"):
        return JSONResponse({"error": "Insufficient permission."}, 403)
    db = get_db()
    try:
        old = db.execute("SELECT * FROM aria_risks WHERE id=%s", (risk_id,)).fetchone()
        if not old:
            return JSONResponse({"error": "Risk not found"}, 404)
        updates, params = [], []
        for field, val in [
            ("status", status), ("owner", owner),
            ("mitigation", mitigation), ("likelihood", likelihood), ("impact", impact),
        ]:
            if val is not None:
                updates.append(field + "=%s")
                params.append(val)
        if updates:
            params.append(risk_id)
            db.execute("UPDATE aria_risks SET " + ", ".join(updates) + " WHERE id=%s", params)
            db.commit()
            log_audit(user, "aria", f"Updated risk #{risk_id} status={status}",
                      "risk", risk_id, old["risk_id"])

            # Emit escalation event when risk status changes to 'Escalated'
            old_status = old["status"] or ""
            if status == "Escalated" and old_status != "Escalated":
                emit(
                    ARIA_RISK_ESCALATED,
                    source_module="aria",
                    entity_type="risk",
                    entity_id=risk_id,
                    payload={
                        "risk_id": old["risk_id"],
                        "description": old["description"],
                        "category": old["category"],
                        "framework": old["framework"],
                        "control_ref": old["control_ref"],
                        "likelihood": likelihood if likelihood is not None else old["likelihood"],
                        "impact": impact if impact is not None else old["impact"],
                        "owner": owner or old["owner"],
                    },
                    user_id=user["id"],
                )
    finally:
        db.close()
    return JSONResponse({"ok": True})


# -- Cross-Framework Mapping (IMS) --------------------------------------------

@router.get("/mapping", response_class=HTMLResponse)
@require_module("aria")
async def mapping_page(request: Request):
    user = request.state.user
    db = get_db()
    try:
        # Active frameworks for the selector
        active_frameworks = db.execute(
            "SELECT id, name, color FROM frameworks WHERE is_active=1 ORDER BY name"
        ).fetchall()
    finally:
        db.close()
    return _aria_render(request, "mapping.html", {
        "user": user, "module": "aria",
        "frameworks": [dict(r) for r in active_frameworks],
        # Legacy static mapping kept for display-only fallback
        "static_mapping": CROSS_MAPPING,
        "framework_names": FRAMEWORK_NAMES,
    }, active_section="mapping")


@router.get("/api/control-mappings")
@require_module("aria")
async def api_list_control_mappings(request: Request):
    """List all control mappings, optionally filtered by framework pair."""
    fw1 = request.query_params.get("fw1")
    fw2 = request.query_params.get("fw2")
    db = get_db()
    try:
        where = ""
        params: list = []
        if fw1 and fw2:
            where = ("WHERE (m.source_framework_id=%s AND m.target_framework_id=%s) "
                     "OR (m.source_framework_id=%s AND m.target_framework_id=%s)")
            params = [fw1, fw2, fw2, fw1]
        elif fw1:
            where = "WHERE m.source_framework_id=%s OR m.target_framework_id=%s"
            params = [fw1, fw1]

        rows = db.execute(f"""
            SELECT m.id, m.mapping_type, m.notes, m.confidence, m.created_at,
                   sf.id AS source_fw_id, sf.name AS source_fw_name,
                   sc.id AS source_ctrl_id, sc.ref AS source_ref, sc.name AS source_name,
                   tf.id AS target_fw_id, tf.name AS target_fw_name,
                   tc.id AS target_ctrl_id, tc.ref AS target_ref, tc.name AS target_name
            FROM aria_control_mappings m
            JOIN frameworks sf ON sf.id = m.source_framework_id
            JOIN controls   sc ON sc.id = m.source_control_id
            JOIN frameworks tf ON tf.id = m.target_framework_id
            JOIN controls   tc ON tc.id = m.target_control_id
            {where}
            ORDER BY sf.name, sc.ref
        """, params).fetchall()
    finally:
        db.close()
    return JSONResponse({"items": [dict(r) for r in rows]})


@router.get("/api/control-mappings/frameworks/{fw_id}/controls")
@require_module("aria")
async def api_framework_controls(request: Request, fw_id: int):
    """Return controls for a given framework (for the mapping UI dropdowns)."""
    db = get_db()
    try:
        rows = db.execute(
            "SELECT id, ref, name, category FROM controls WHERE framework_id=%s ORDER BY ref",
            (fw_id,)
        ).fetchall()
    finally:
        db.close()
    return JSONResponse({"items": [dict(r) for r in rows]})


@router.post("/api/control-mappings")
@require_module("aria")
async def api_create_control_mapping(request: Request):
    """Create a new cross-framework control mapping."""
    user = request.state.user
    body = await request.json()
    source_ctrl_id = body.get("source_control_id")
    target_ctrl_id = body.get("target_control_id")
    mapping_type   = body.get("mapping_type", "equivalent")
    notes          = body.get("notes", "")

    if not source_ctrl_id or not target_ctrl_id:
        return JSONResponse({"ok": False, "error": "source_control_id and target_control_id are required"}, status_code=400)
    if source_ctrl_id == target_ctrl_id:
        return JSONResponse({"ok": False, "error": "Cannot map a control to itself"}, status_code=400)

    db = get_db()
    try:
        # Look up framework IDs from control IDs
        src = db.execute("SELECT id, framework_id FROM controls WHERE id=%s", (source_ctrl_id,)).fetchone()
        tgt = db.execute("SELECT id, framework_id FROM controls WHERE id=%s", (target_ctrl_id,)).fetchone()
        if not src or not tgt:
            return JSONResponse({"ok": False, "error": "One or both controls not found"}, status_code=404)
        if src["framework_id"] == tgt["framework_id"]:
            return JSONResponse({"ok": False, "error": "Both controls belong to the same framework — cross-framework mapping only"}, status_code=400)

        new_id = insert_returning_id(db,"""
            INSERT INTO aria_control_mappings
            (source_framework_id, source_control_id, target_framework_id, target_control_id,
             mapping_type, notes, created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING
        """, (src["framework_id"], source_ctrl_id,
              tgt["framework_id"], target_ctrl_id,
              mapping_type, notes, user["id"]))
        db.commit()
        if new_id is None:
            # ON CONFLICT DO NOTHING fired — mapping already exists
            existing = db.execute(
                "SELECT id FROM aria_control_mappings WHERE source_control_id=%s AND target_control_id=%s",
                (source_ctrl_id, target_ctrl_id)
            ).fetchone()
            return JSONResponse({"ok": True, "id": existing["id"] if existing else None, "already_existed": True})
    finally:
        db.close()
    return JSONResponse({"ok": True, "id": new_id})


@router.delete("/api/control-mappings/{mapping_id}")
@require_module("aria")
async def api_delete_control_mapping(request: Request, mapping_id: int):
    """Delete a control mapping by ID."""
    db = get_db()
    try:
        db.execute("DELETE FROM aria_control_mappings WHERE id=%s", (mapping_id,))
        db.commit()
    finally:
        db.close()
    return JSONResponse({"ok": True})


@router.post("/api/control-mappings/auto-generate")
@require_module("aria")
async def api_auto_generate_mappings(request: Request):
    """Automatically generate control mappings across selected frameworks.

    Body: {"framework_ids": [1, 4, 7], "use_ai": true}
    Returns: {"ok": true, "created": N, "skipped": M, "ai_calls": K}
    """
    try:
        user = request.state.user
        try:
            body = await request.json()
        except Exception:
            return JSONResponse({"ok": False, "error": "Invalid request body"}, status_code=400)
        fw_ids = body.get("framework_ids", [])
        use_ai = bool(body.get("use_ai", False))

        if len(fw_ids) < 2:
            return JSONResponse({"ok": False, "error": "Select at least 2 frameworks"}, status_code=400)
        if len(fw_ids) > 6:
            return JSONResponse({"ok": False, "error": "Maximum 6 frameworks per run"}, status_code=400)

        from core.auto_mapper import run_auto_mapping
        db = get_db()
        try:
            result = await run_auto_mapping(
                framework_ids=[int(f) for f in fw_ids],
                user_id=user["id"],
                db=db,
                use_ai=use_ai,
            )
        finally:
            db.close()
        return JSONResponse({"ok": True, **result})
    except Exception as exc:
        log.exception("auto-generate failed: %s", exc)
        return JSONResponse({"ok": False, "error": str(exc)}, status_code=500)


@router.get("/api/control-mappings/ims-status")
@require_module("aria")
async def api_ims_status(request: Request):
    """Return IMS classification for all controls across the given frameworks.

    Query param: fw_ids=1,2,3
    Returns three sections: integrated, partial, unique — each a list of control dicts.
    """
    _empty = {"integrated": [], "partial": [], "unique": [], "fw_ids": [], "stats": {
        "total": 0, "integrated": 0, "partial": 0, "unique": 0, "effort_saved_pct": 0}}
    try:
        fw_ids_param = request.query_params.get("fw_ids", "")
        try:
            fw_ids = [int(x) for x in fw_ids_param.split(",") if x.strip().isdigit()]
        except ValueError:
            fw_ids = []

        if len(fw_ids) < 2:
            return JSONResponse({**_empty, "fw_ids": fw_ids})

        # Pre-seed curated mappings from official standard annexes
        try:
            from seeds.framework_controls import seed_curated_mappings
            seed_curated_mappings()
        except Exception:
            pass

        db = get_db()
        try:
            placeholders = ",".join(["%s"] * len(fw_ids))
            rows = db.execute(
                f"SELECT c.id, c.framework_id, c.ref, c.name, c.description, c.category, "
                f"       f.name AS fw_name, f.color AS fw_color "
                f"FROM controls c JOIN frameworks f ON f.id=c.framework_id "
                f"WHERE c.framework_id IN ({placeholders}) ORDER BY f.name, c.ref",
                fw_ids,
            ).fetchall()
            controls = [dict(r) for r in rows]

            from core.auto_mapper import get_ims_status_bulk
            statuses = get_ims_status_bulk(controls, fw_ids, db)

            doc_refs = set()
            try:
                doc_rows = db.execute(
                    "SELECT DISTINCT framework, control_ref FROM aria_documents"
                ).fetchall()
                for dr in doc_rows:
                    doc_refs.add((dr["framework"], dr["control_ref"]))
            except Exception:
                pass
        finally:
            db.close()

        integrated, partial, unique = [], [], []
        for ctrl in controls:
            st = statuses.get(ctrl["id"], {})
            ctrl["ims_status"]      = st.get("status", "unique")
            ctrl["mapped_fw_ids"]   = st.get("mapped_fw_ids", [])
            ctrl["mapped_fw_names"] = st.get("mapped_fw_names", [])
            ctrl["has_doc"] = (ctrl.get("fw_name", ""), ctrl.get("ref", "")) in doc_refs
            if ctrl["ims_status"] == "integrated":
                integrated.append(ctrl)
            elif ctrl["ims_status"] == "partial":
                partial.append(ctrl)
            else:
                unique.append(ctrl)

        return JSONResponse({
            "integrated": integrated,
            "partial":    partial,
            "unique":     unique,
            "fw_ids":     fw_ids,
            "stats": {
                "total":      len(controls),
                "integrated": len(integrated),
                "partial":    len(partial),
                "unique":     len(unique),
                "effort_saved_pct": (
                    round((len(integrated) * 100) / len(controls) / 2) if controls else 0
                ),
            },
        })
    except Exception as exc:
        log.exception("ims-status failed: %s", exc)
        return JSONResponse({**_empty, "error": str(exc)}, status_code=500)


# -- Excel Export -------------------------------------------------------------

@router.get("/export/excel")
@require_module("aria")
async def export_excel(request: Request):
    user = request.state.user
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    db = get_db()
    try:
        frameworks = db.execute(
            "SELECT * FROM frameworks WHERE is_active = 1 ORDER BY name"
        ).fetchall()
        wb = Workbook()
        wb.remove(wb.active)
        hdr_fill = PatternFill("solid", start_color="1A2744", end_color="1A2744")
        hdr_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for fw in frameworks:
            ws = wb.create_sheet(fw["name"][:31])
            ws.sheet_view.showGridLines = False
            headers = [
                "Ref", "Control Name", "Category", "Doc Type",
                "Document Title", "Owner", "Status", "Priority",
                "Target Date", "Review Date", "Version", "Notes",
            ]
            for c, h in enumerate(headers, 1):
                cell = ws.cell(row=1, column=c, value=h)
                cell.font = hdr_font
                cell.fill = hdr_fill
                cell.alignment = Alignment(
                    horizontal="center", vertical="center"
                )
                cell.border = border
            ws.row_dimensions[1].height = 24

            controls = db.execute(
                "SELECT * FROM controls WHERE framework_id=%s ORDER BY ref",
                (fw["id"],),
            ).fetchall()
            for i, ctrl in enumerate(controls, 2):
                bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
                row_fill = PatternFill("solid", start_color=bg, end_color=bg)
                vals = [
                    ctrl["ref"], ctrl["name"], ctrl["category"],
                    ctrl["doc_type"], ctrl["document_title"], ctrl["owner"],
                    ctrl["status"], ctrl["priority"], ctrl["target_date"],
                    ctrl["review_date"], ctrl["version"], ctrl["notes"],
                ]
                for c, v in enumerate(vals, 1):
                    cell = ws.cell(row=i, column=c, value=v or "")
                    cell.fill = row_fill
                    cell.border = border
                    cell.font = Font(name="Calibri", size=9)
                    cell.alignment = Alignment(
                        vertical="center", wrap_text=c in (2, 3, 5, 12)
                    )
                ws.row_dimensions[i].height = 18

            widths = [10, 28, 22, 12, 35, 18, 16, 12, 14, 14, 10, 30]
            for c, w in enumerate(widths, 1):
                col_letter = ws.cell(row=1, column=c).column_letter
                ws.column_dimensions[col_letter].width = w
    finally:
        db.close()

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = "ARIA_Export_" + datetime.now().strftime("%Y%m%d_%H%M") + ".xlsx"
    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".spreadsheetml.sheet"
        ),
        headers={"Content-Disposition": "attachment; filename=" + fname},
    )


# -- Audit Log (ARIA-scoped) -------------------------------------------------

@router.get("/audit-log", response_class=HTMLResponse)
@require_module("aria")
async def audit_log_page(request: Request):
    user = request.state.user
    db = get_db()
    try:
        logs = db.execute("""
            SELECT al.*, u.full_name FROM audit_log al
            LEFT JOIN users u ON al.user_id = u.id
            WHERE al.module = 'aria'
            ORDER BY al.created_at DESC LIMIT 200
        """).fetchall()
    finally:
        db.close()
    return _aria_render(request, "audit_log.html", {
        "user": user, "module": "aria", "logs": logs,
    }, active_section="audit-log")


# -- AI Generator -------------------------------------------------------------

@router.get("/ai-generator", response_class=HTMLResponse)
@require_module("aria")
async def ai_generator_page(request: Request):
    user = request.state.user
    db = get_db()
    try:
        frameworks = db.execute(
            "SELECT * FROM frameworks WHERE is_active = 1 ORDER BY name"
        ).fetchall()
        log.info("AI-Generator: %d active frameworks found", len(frameworks))

        # Import seed data once for this request
        try:
            from seeds.framework_controls import FRAMEWORK_CONTROLS
        except Exception as imp_err:
            log.exception("Cannot import FRAMEWORK_CONTROLS: %s", imp_err)
            FRAMEWORK_CONTROLS = {}

        # Pre-load ALL controls for active frameworks
        controls_by_fw = {}
        for fw in frameworks:
            fw_id = fw["id"]
            fw_name = fw["name"]

            ctrls = db.execute(
                "SELECT id, ref, name, description, doc_type, status "
                "FROM controls WHERE framework_id=%s ORDER BY ref",
                (fw_id,),
            ).fetchall()
            log.info("  Framework '%s' (id=%s): %d controls",
                     fw_name, fw_id, len(ctrls))

            # Lazy-seed if empty — try exact match, then partial match
            if not ctrls:
                seed_data = FRAMEWORK_CONTROLS.get(fw_name)
                if not seed_data:
                    # Try partial/case-insensitive match
                    fw_lower = fw_name.lower()
                    for key, data in FRAMEWORK_CONTROLS.items():
                        if (key.lower() in fw_lower
                                or fw_lower in key.lower()):
                            seed_data = data
                            log.info("  Partial match: '%s' → '%s'",
                                     fw_name, key)
                            break

                if seed_data:
                    log.info("  Lazy-seeding %d controls for '%s'",
                             len(seed_data), fw_name)
                    inserted = 0
                    for ctrl in seed_data:
                        try:
                            db.execute(
                                "INSERT INTO controls "
                                "(framework_id, ref, name, description, "
                                "category, doc_type, priority, last_updated)"
                                " VALUES (%s, %s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING",
                                (
                                    fw_id,
                                    ctrl.get("ref", ""),
                                    ctrl.get("name", ""),
                                    ctrl.get("description", ""),
                                    ctrl.get("category", ""),
                                    ctrl.get("doc_type", "Policy"),
                                    ctrl.get("priority", "High"),
                                    utcnow().isoformat(),
                                ),
                            )
                            inserted += 1
                        except Exception as ins_err:
                            log.warning("  Insert failed ref=%s: %s",
                                        ctrl.get("ref"), ins_err)
                    db.commit()
                    log.info("  Inserted %d controls", inserted)

                    # Update framework total
                    db.execute(
                        "UPDATE frameworks SET total_controls = "
                        "(SELECT COUNT(*) FROM controls "
                        "WHERE framework_id = %s) WHERE id = %s",
                        (fw_id, fw_id),
                    )
                    db.commit()

                    # Re-fetch
                    ctrls = db.execute(
                        "SELECT id, ref, name, description, "
                        "doc_type, status "
                        "FROM controls WHERE framework_id=%s "
                        "ORDER BY ref",
                        (fw_id,),
                    ).fetchall()
                    log.info("  After seeding: %d controls", len(ctrls))
                else:
                    log.warning("  No seed data found for '%s'. "
                                "Available keys: %s",
                                fw_name, list(FRAMEWORK_CONTROLS.keys()))

            controls_by_fw[str(fw_id)] = [dict(c) for c in ctrls]

        total_controls = sum(len(v) for v in controls_by_fw.values())
        log.info("AI-Generator: total %d controls across %d frameworks",
                 total_controls, len(controls_by_fw))
    finally:
        db.close()

    # ── Document stats & recent docs ──────────────────────────────────────────
    db3 = get_db()
    try:
        doc_rows = db3.execute(
            "SELECT framework, COUNT(*) as cnt FROM aria_documents GROUP BY framework"
        ).fetchall()
        docs_by_fw_map = {r["framework"]: r["cnt"] for r in doc_rows}
        total_docs = sum(docs_by_fw_map.values())

        controls_with_docs = db3.execute(
            "SELECT COUNT(DISTINCT control_ref || '|' || framework) FROM aria_documents"
        ).fetchone()[0]

        recent_docs = [dict(r) for r in db3.execute(
            "SELECT doc_id, title, framework, doc_type, control_ref, updated_at "
            "FROM aria_documents ORDER BY updated_at DESC LIMIT 5"
        ).fetchall()]

        doc_ctrl_ref_set = [
            r["framework"] + "|" + r["control_ref"]
            for r in db3.execute(
                "SELECT DISTINCT framework, control_ref FROM aria_documents"
            ).fetchall()
        ]
    finally:
        db3.close()

    # ── AI provider label ──────────────────────────────────────────────────────
    _provider_key = (settings.AI_PROVIDER or "anthropic").lower()
    _ollama_model = getattr(settings, "OLLAMA_MODEL", "llama3.2")
    ai_provider_label = {
        "anthropic":       "Claude (Anthropic)",
        "openai":          f"GPT-4o (OpenAI)",
        "deepseek":        "DeepSeek Chat",
        "ollama":          f"Ollama · {_ollama_model}",
        "gemini":          "Gemini (Google)",
    }.get(_provider_key, "Claude (Anthropic)")

    api_configured = bool(
        settings.ANTHROPIC_API_KEY or
        getattr(settings, "OPENAI_API_KEY", "") or
        getattr(settings, "DEEPSEEK_API_KEY", "") or
        _provider_key == "ollama"
    )

    return _aria_render(request, "ai_generator.html", {
        "user": user, "module": "aria",
        "frameworks": frameworks,
        "active_frameworks": frameworks,  # alias used by IMS dropdown
        "api_configured": api_configured,
        "controls_by_fw_json":    json.dumps(controls_by_fw, ensure_ascii=True),
        "docs_by_fw_map":         docs_by_fw_map,
        "total_docs":             total_docs,
        "total_controls":         total_controls,
        "controls_with_docs":     controls_with_docs,
        "recent_docs":            recent_docs,
        "doc_ctrl_refs_json":     json.dumps(doc_ctrl_ref_set, ensure_ascii=True),
        "ai_provider_label":      ai_provider_label,
    }, active_section="ai-generator")


@router.post("/api/generate-policy")
@require_module("aria")
async def api_generate_policy(request: Request,
                              control_id: int = Form(...),
                              org_name: str = Form("Your Organisation"),
                              doc_type_override: str = Form(""),
                              integrated_framework_id: str = Form("")):
    """Generate a governance document for a control.

    Set integrated_framework_id to a framework ID (or comma-separated IDs) to
    generate an Integrated Management System (IMS) document that covers the
    primary control AND all mapped controls from the specified additional framework(s).
    """
    user = request.state.user
    if not has_capability(user, "aria.policy.generate_ai"):
        return JSONResponse({
            "error": "You need Policy Author or Compliance Manager role."
        }, 403)

    db = get_db()
    try:
        ctrl = db.execute("""
            SELECT c.*, f.name as fw_name
            FROM controls c
            JOIN frameworks f ON c.framework_id = f.id
            WHERE c.id=%s
        """, (control_id,)).fetchone()

        # IMS: resolve additional mapped controls for extra frameworks
        integrated_frameworks = []
        if integrated_framework_id:
            target_fw_ids = [int(x.strip()) for x in integrated_framework_id.split(",") if x.strip().isdigit()]
            for tfw_id in target_fw_ids:
                mapped = db.execute("""
                    SELECT c.ref, c.name, c.description, f.name AS fw_name
                    FROM aria_control_mappings m
                    JOIN controls c   ON c.id = m.target_control_id
                    JOIN frameworks f ON f.id = m.target_framework_id
                    WHERE m.source_control_id = %s AND m.target_framework_id = %s
                    UNION
                    SELECT c.ref, c.name, c.description, f.name AS fw_name
                    FROM aria_control_mappings m
                    JOIN controls c   ON c.id = m.source_control_id
                    JOIN frameworks f ON f.id = m.source_framework_id
                    WHERE m.target_control_id = %s AND m.source_framework_id = %s
                """, (control_id, tfw_id, control_id, tfw_id)).fetchall()
                for row in mapped:
                    integrated_frameworks.append({
                        "framework": row["fw_name"],
                        "ref": row["ref"],
                        "name": row["name"],
                        "description": row["description"] or "",
                    })
    finally:
        db.close()

    if not ctrl:
        return JSONResponse({"error": "Control not found"}, 404)

    # Honour user's doc-type override if provided
    resolved_doc_type = (doc_type_override or "").strip() or ctrl["doc_type"]

    from modules.aria.ai_generator import generate_policy
    result = await generate_policy(
        framework=ctrl["fw_name"],
        control_ref=ctrl["ref"],
        control_name=ctrl["name"],
        control_description=ctrl["description"],
        doc_type=resolved_doc_type,
        org_name=org_name,
        integrated_frameworks=integrated_frameworks if integrated_frameworks else None,
    )

    if result["success"]:
        log_audit(user, "aria",
                  "Generated AI policy for " + ctrl["fw_name"] + " " + ctrl["ref"],
                  "control", control_id)
        db = get_db()
        try:
            db.execute(
                "UPDATE controls SET last_updated=%s WHERE id=%s",
                (datetime.now().strftime("%Y-%m-%d"), control_id),
            )
            existing = db.execute(
                "SELECT doc_id FROM aria_documents "
                "WHERE control_ref=%s AND framework=%s",
                (ctrl["ref"], ctrl["fw_name"]),
            ).fetchone()
            now = datetime.now().isoformat()
            policy_body = result.get("content", "")
            if existing:
                old_ver = db.execute(
                    "SELECT version FROM aria_documents WHERE doc_id=%s",
                    (existing["doc_id"],),
                ).fetchone()["version"]
                try:
                    parts = old_ver.split(".")
                    new_ver = parts[0] + "." + str(int(parts[1]) + 1)
                except (IndexError, ValueError):
                    new_ver = "1.1"
                db.execute("""
                    UPDATE aria_documents
                    SET version=%s, status='Draft', updated_at=%s,
                        comments='AI Generated -- updated ' || %s,
                        body=%s
                    WHERE doc_id=%s
                """, (new_ver, now, datetime.now().strftime("%Y-%m-%d"),
                      policy_body, existing["doc_id"]))
            else:
                count = db.execute(
                    "SELECT COUNT(*) FROM aria_documents"
                ).fetchone()[0]
                doc_id = "DOC-%04d" % (count + 1)
                db.execute("""
                    INSERT INTO aria_documents
                    (doc_id, framework, control_ref, title, doc_type,
                     version, status, owner, created_at, updated_at,
                     comments, body)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """, (doc_id, ctrl["fw_name"], ctrl["ref"],
                      ctrl["name"] + " -- " + resolved_doc_type,
                      resolved_doc_type, "1.0", "Draft",
                      user.get("full_name", ""),
                      now, now,
                      "AI Generated on " + datetime.now().strftime("%Y-%m-%d"),
                      policy_body))
            db.commit()
        finally:
            db.close()

    return JSONResponse(result)


@router.post("/api/generate-gap-analysis")
@require_module("aria")
async def api_gap_analysis(request: Request,
                           framework_id: int = Form(...)):
    user = request.state.user
    if not has_capability(user, "aria.policy.generate_ai"):
        return JSONResponse({
            "error": "You need Policy Author or Compliance Manager role."
        }, 403)

    db = get_db()
    try:
        fw = db.execute(
            "SELECT * FROM frameworks WHERE id=%s", (framework_id,)
        ).fetchone()
        controls = db.execute(
            "SELECT ref, name, status, priority "
            "FROM controls WHERE framework_id=%s",
            (framework_id,),
        ).fetchall()
    finally:
        db.close()

    if not fw:
        return JSONResponse({"error": "Framework not found"}, 404)

    from modules.aria.ai_generator import generate_gap_analysis
    result = await generate_gap_analysis(
        fw["name"], [dict(c) for c in controls]
    )
    if result["success"]:
        log_audit(user, "aria",
                  "Generated gap analysis for " + fw["name"],
                  "framework", framework_id)
    return JSONResponse(result)


@router.post("/api/export-word")
@require_module("aria")
async def export_word(request: Request, control_id: str = Form(""), content: str = Form(...), org_name: str = Form("Your Organisation")):
    """Convert AI-generated policy markdown to a .docx download."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    db = get_db()
    try:
        ctrl_row = db.execute("SELECT ref, title FROM controls WHERE id=%s", (control_id,)).fetchone() if control_id else None
    finally:
        db.close()

    ctrl_label = f"{ctrl_row['ref']} - {ctrl_row['title']}" if ctrl_row else "Policy Document"

    heading = doc.add_heading(f"{org_name}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading(ctrl_label, level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")

    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_ref = (ctrl_row["ref"].replace(".", "_") if ctrl_row else "policy")
    filename = f"ARIA_{safe_ref}_Policy.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# -- Ask ARIA -----------------------------------------------------------------

@router.get("/ask", response_class=HTMLResponse)
@require_module("aria")
async def ask_page(request: Request):
    user = request.state.user
    db = get_db()
    try:
        recent = db.execute("""
            SELECT question, answer, covered, citations, created_at
            FROM aria_ask_log
            WHERE username=%s
            ORDER BY id DESC LIMIT 12
        """, (user["username"],)).fetchall()
        suggestions_rows = db.execute("""
            SELECT title FROM aria_documents
            WHERE body IS NOT NULL AND length(body) > 200
            ORDER BY RANDOM() LIMIT 5
        """).fetchall()
    finally:
        db.close()

    suggestions = [
        "What does our " + r["title"] + " say?" for r in suggestions_rows
    ] or [
        "What is our remote working policy?",
        "How do I report a security incident?",
        "Can I use ChatGPT for work?",
        "What counts as confidential information?",
        "How long do we keep customer data?",
    ]

    # ── Stats for the header strip ─────────────────────────────────────────
    db2 = get_db()
    try:
        total_indexed = db2.execute(
            "SELECT COUNT(*) FROM aria_ask_index"
        ).fetchone()[0]
        total_asked = db2.execute(
            "SELECT COUNT(*) FROM aria_ask_log WHERE username=%s",
            (user["username"],)
        ).fetchone()[0]
        covered_row = db2.execute(
            "SELECT ROUND(AVG(CASE WHEN covered=1 THEN 100.0 ELSE 0 END),0) "
            "FROM aria_ask_log WHERE username=%s",
            (user["username"],)
        ).fetchone()[0]
        covered_pct = int(covered_row) if covered_row is not None else 0
        helpful_row = db2.execute(
            "SELECT ROUND(AVG(CASE WHEN feedback=1 THEN 100.0 ELSE 0 END),0) "
            "FROM aria_ask_log WHERE username=%s AND feedback IS NOT NULL",
            (user["username"],)
        ).fetchone()[0]
        helpful_pct = int(helpful_row) if helpful_row is not None else None

        frameworks_list = db2.execute(
            "SELECT name FROM frameworks WHERE is_active=1 ORDER BY name"
        ).fetchall()
    except Exception:
        total_indexed = 0; total_asked = 0; covered_pct = 0
        helpful_pct = None; frameworks_list = []
    finally:
        db2.close()

    from config import settings as _cfg
    _provider_key = (getattr(_cfg, 'AI_PROVIDER', '') or 'anthropic').lower()
    _ollama_m = getattr(_cfg, 'OLLAMA_MODEL', 'llama3.2')
    ai_provider_label = {
        "anthropic": "Claude",
        "openai":    "GPT-4o",
        "deepseek":  "DeepSeek",
        "ollama":    f"Ollama · {_ollama_m}",
        "gemini":    "Gemini",
    }.get(_provider_key, "Claude")

    return _aria_render(request, "ask.html", {
        "user": user, "module": "aria",
        "recent": recent, "suggestions": suggestions,
        "total_indexed":    total_indexed,
        "total_asked":      total_asked,
        "covered_pct":      covered_pct,
        "helpful_pct":      helpful_pct,
        "frameworks_list":  [r["name"] for r in frameworks_list],
        "ai_provider_label": ai_provider_label,
    }, active_section="ask")


@router.post("/api/ask")
@require_module("aria")
async def api_ask(request: Request,
                  question: str = Form(...),
                  framework_filter: str = Form("")):
    user = request.state.user
    question = (question or "").strip()
    if not question:
        return JSONResponse({"error": "Empty question"}, 400)
    if len(question) > 2000:
        return JSONResponse({"error": "Question too long"}, 400)

    from modules.aria.ask_service import ask as ask_policy
    result = await ask_policy(question, user=user,
                               framework_filter=(framework_filter or "").strip())
    log_audit(user, "aria", "Asked ARIA: " + question[:80], "ask")
    return JSONResponse(result)


@router.post("/api/ask/feedback")
@require_module("aria")
async def api_ask_feedback(request: Request,
                            log_id: int = Form(...),
                            feedback: int = Form(...)):
    """Store thumbs-up (1) or thumbs-down (-1) on a Q&A log entry."""
    if feedback not in (1, -1):
        return JSONResponse({"error": "feedback must be 1 or -1"}, 400)
    user = request.state.user
    db = get_db()
    try:
        db.execute(
            "UPDATE aria_ask_log SET feedback=%s WHERE id=%s AND user_id=%s",
            (feedback, log_id, user["id"]),
        )
        db.commit()
    finally:
        db.close()
    return JSONResponse({"ok": True})


@router.post("/api/ask/rebuild")
@require_module("aria")
async def api_ask_rebuild(request: Request):
    user = request.state.user
    if not has_capability(user, "platform.manage_users"):
        return JSONResponse({"error": "Admin only"}, 403)
    from modules.aria.ask_service import rebuild_all
    try:
        n = rebuild_all()
        log_audit(user, "aria",
                  "Rebuilt Ask ARIA search index (" + str(n) + " chunks)",
                  "ask_index")
        return JSONResponse({"ok": True, "chunks_indexed": n})
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)
